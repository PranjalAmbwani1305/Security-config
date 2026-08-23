"""
Security Configuration & CIS Benchmark Compliance Tool
--------------------------------------------------------
A Streamlit app that lets a consultant/auditor walk a client through a
hardening checklist for a chosen technology (based on CIS Benchmark
categories and general security best practice where a formal CIS
benchmark doesn't apply), record compliance status + evidence/notes per
item, and view a live compliance dashboard.

Run with:
    pip install streamlit pandas plotly gspread google-auth
    streamlit run security_config_dashboard.py

Cloud persistence (optional):
    Set up a Google Sheet + service account and add its credentials to
    Streamlit secrets (see README.md -> "Cloud persistence setup") to get
    automatic save/load per client that survives app restarts on Streamlit
    Community Cloud. Without secrets configured, the app still works fully
    using the CSV export/import already built in.
"""

import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
import io
import json

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import gspread
    from google.oauth2.service_account import Credentials
    GSPREAD_AVAILABLE = True
except ImportError:
    GSPREAD_AVAILABLE = False

SHEETS_SCOPE = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]


@st.cache_resource(show_spinner=False)
def get_gsheet_client():
    """Build an authorized gspread client from Streamlit secrets, if configured."""
    if not GSPREAD_AVAILABLE:
        return None
    if "gcp_service_account" not in st.secrets or "sheet_id" not in st.secrets.get("gsheets", {}):
        return None
    creds = Credentials.from_service_account_info(
        dict(st.secrets["gcp_service_account"]), scopes=SHEETS_SCOPE
    )
    return gspread.authorize(creds)


def cloud_persistence_enabled() -> bool:
    return get_gsheet_client() is not None


def _sheet_tab_name(client_name: str, technology: str) -> str:
    # Sheet tab names are capped at 100 chars and can't contain some symbols.
    raw = f"{client_name}__{technology}".strip() or "untitled"
    safe = "".join(c for c in raw if c.isalnum() or c in ("_", "-", " "))
    return safe[:95]


def save_to_cloud(client_name: str, technology: str, responses: dict) -> None:
    gc = get_gsheet_client()
    if gc is None:
        return
    sh = gc.open_by_key(st.secrets["gsheets"]["sheet_id"])
    tab_name = _sheet_tab_name(client_name, technology)
    try:
        ws = sh.worksheet(tab_name)
        ws.clear()
    except gspread.WorksheetNotFound:
        ws = sh.add_worksheet(title=tab_name, rows=200, cols=10)

    rows = [["item_id", "category", "control", "reference", "audit_step", "status", "notes", "reviewer", "date"]]
    for item_id, data in responses.items():
        rows.append([
            item_id,
            data.get("category", ""),
            data.get("control", ""),
            data.get("reference", ""),
            data.get("audit_step", ""),
            data.get("status", ""),
            data.get("notes", ""),
            data.get("reviewer", ""),
            data.get("date", ""),
        ])
    ws.update(rows)


def load_from_cloud(client_name: str, technology: str) -> dict:
    gc = get_gsheet_client()
    if gc is None:
        return {}
    sh = gc.open_by_key(st.secrets["gsheets"]["sheet_id"])
    tab_name = _sheet_tab_name(client_name, technology)
    try:
        ws = sh.worksheet(tab_name)
    except gspread.WorksheetNotFound:
        return {}

    records = ws.get_all_records()
    loaded = {}
    for row in records:
        item_id = row.get("item_id")
        if not item_id:
            continue
        loaded[item_id] = {
            "status": row.get("status", "Not Reviewed") or "Not Reviewed",
            "notes": row.get("notes", ""),
            "category": row.get("category", ""),
            "control": row.get("control", ""),
            "reference": row.get("reference", ""),
            "audit_step": row.get("audit_step", ""),
            "reviewer": row.get("reviewer", ""),
            "date": row.get("date", ""),
        }
    return loaded


# ---------------------------------------------------------------------------
# HUGGING FACE - AI-assisted features
# Uses huggingface_hub.InferenceClient against a chat-capable instruct model
# routed through HF's Inference Providers (Together AI, SambaNova, Cerebras,
# Groq, Novita, etc). NOTE: HF's free serverless API for small models
# (e.g. Qwen2.5-7B-Instruct, Llama-3.1-8B-Instruct) has largely been
# deprecated in favor of Inference Providers, which only host a curated set
# of models per partner. Pick a default that is actually still routed by at
# least one provider - see https://huggingface.co/models?inference_provider=all
# Configure via Streamlit secrets:
#   [huggingface]
#   api_key = "hf_xxx"
#   model = "meta-llama/Llama-3.3-70B-Instruct"   # optional, has a default
# ---------------------------------------------------------------------------

try:
    from huggingface_hub import InferenceClient
    HF_AVAILABLE = True
except ImportError:
    HF_AVAILABLE = False

# Verified LIVE against https://huggingface.co/models?inference_provider=all&pipeline_tag=text-generation
# at the time this was last checked. Provider coverage rotates over time
# (models get dropped with no warning, which is what caused this whole
# fallback chain to go stale twice already) - if this stops working, don't
# guess a replacement, re-check that URL for a model that is CURRENTLY
# filtered as having an inference provider before swapping it in.
DEFAULT_HF_MODEL = "openai/gpt-oss-120b"

# Alternates confirmed live on the same check, used only to populate the
# error message below.
_FALLBACK_MODEL_SUGGESTIONS = [
    "openai/gpt-oss-120b",
    "meta-llama/Llama-3.1-8B-Instruct",
    "deepseek-ai/DeepSeek-R1",
]


def hf_diagnose() -> str:
    """Returns None if everything needed for AI features is in place, otherwise
    a specific human-readable reason. Checked independently of client creation
    so a bad token doesn't get lumped in with 'not configured' at all."""
    if not HF_AVAILABLE:
        return "The `huggingface_hub` package isn't installed. Add `huggingface_hub` to requirements.txt and reboot the app."
    if "huggingface" not in st.secrets:
        return "No `[huggingface]` section found in Streamlit secrets."
    hf_secrets = st.secrets.get("huggingface", {})
    if "api_key" not in hf_secrets or not str(hf_secrets.get("api_key", "")).strip():
        return "`[huggingface]` section exists in secrets but `api_key` is missing or empty."
    token = str(hf_secrets["api_key"]).strip()
    if not token.startswith("hf_"):
        return "The configured api_key doesn't look like a Hugging Face token (should start with 'hf_')."
    return None


@st.cache_resource(show_spinner=False)
def get_hf_client():
    if hf_diagnose() is not None:
        return None
    model = st.secrets["huggingface"].get("model", DEFAULT_HF_MODEL)
    try:
        return InferenceClient(model=model, token=st.secrets["huggingface"]["api_key"], provider="auto")
    except Exception:
        return None


def ai_enabled() -> bool:
    return get_hf_client() is not None


def ask_ai(prompt: str, system: str = "", max_tokens: int = 900) -> str:
    """Send a chat completion request to the configured HF model. Returns the
    text response, or an error string prefixed with '⚠️' on failure — callers
    should check for that prefix before treating the result as trustworthy."""
    client = get_hf_client()
    if client is None:
        reason = hf_diagnose() or "the client failed to initialize for an unknown reason"
        return f"⚠️ AI features are not available: {reason}"
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})
    try:
        resp = client.chat_completion(messages=messages, max_tokens=max_tokens, temperature=0.3)
        return resp.choices[0].message.content
    except Exception as e:
        err_text = str(e)
        if "not supported by any provider" in err_text or "model_not_supported" in err_text:
            configured_model = st.secrets.get("huggingface", {}).get("model", DEFAULT_HF_MODEL)
            # Don't recommend the same dead model back to the user - only
            # suggest alternates that differ from whatever just failed.
            alternates = [m for m in _FALLBACK_MODEL_SUGGESTIONS if m != configured_model]
            if not alternates:
                alternates = _FALLBACK_MODEL_SUGGESTIONS
            suggestion_list = ", ".join(f'"{m}"' for m in alternates)
            return (
                f"⚠️ The model '{configured_model}' is no longer served by any Hugging Face "
                f"Inference Provider. Try setting model to one of: {suggestion_list} "
                f"under [huggingface] in secrets, or check "
                f"https://huggingface.co/models?inference_provider=all&pipeline_tag=text-generation "
                f"for currently-served models."
            )
        return f"⚠️ AI request failed: {e}"


# ---------------------------------------------------------------------------
# AGENTIC LAYER
# This is what actually makes checklist drafting "agentic" rather than a
# single chatbot call: the agent validates its own output, and on failure
# feeds its own error back to itself and retries — without asking the human
# to fix anything — before finally handing back a result (or an honest
# failure after exhausting its attempts). ask_ai() alone has none of this;
# it is one request, one response, no self-checking.
# ---------------------------------------------------------------------------

def agentic_json_call(prompt: str, system: str, max_tokens: int = 1600, max_attempts: int = 3):
    """
    Calls the model expecting a JSON object/array back. If the response
    doesn't parse, the agent appends its own parse error to the prompt and
    retries automatically, up to max_attempts total. Returns
    (parsed_json_or_None, attempts_used, last_raw_response).
    """
    current_prompt = prompt
    last_raw = ""
    for attempt in range(1, max_attempts + 1):
        raw = ask_ai(current_prompt, system=system, max_tokens=max_tokens)
        last_raw = raw
        if raw.startswith("⚠️"):
            return None, attempt, raw  # infrastructure failure, not a JSON failure — don't retry

        cleaned = raw.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.strip("`")
            cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned
            if cleaned.rstrip().endswith("json"):
                cleaned = cleaned.rstrip()[:-4]

        try:
            parsed = json.loads(cleaned)
            return parsed, attempt, raw
        except json.JSONDecodeError as e:
            if attempt == max_attempts:
                return None, attempt, raw
            # Self-correction step: the agent sees its own mistake and retries.
            current_prompt = (
                f"{prompt}\n\n"
                f"Your previous response could not be parsed as JSON. "
                f"Parser error: {e}\n"
                f"Your previous response was:\n{raw[:800]}\n\n"
                f"Return ONLY valid JSON this time — no markdown fences, no commentary, "
                f"no text before or after the JSON."
            )
    return None, max_attempts, last_raw


st.set_page_config(page_title="Sentinel GRC | Security Configuration & Compliance Platform", layout="wide", page_icon="🛡️")

# ---------------------------------------------------------------------------
# THEME - dark, professional GRC-console styling
# ---------------------------------------------------------------------------

st.markdown("""
<style>
#MainMenu, footer, header {visibility: hidden;}
.stApp {
    background: linear-gradient(180deg, #0b1120 0%, #0f172a 100%);
}
h1, h2, h3, h4 { color: #e2e8f0 !important; letter-spacing: -0.01em; }
p, li, span, label, .stMarkdown { color: #cbd5e1; }
.stCaption, [data-testid="stCaptionContainer"] { color: #64748b !important; }

.hero-banner {
    background: linear-gradient(120deg, #0f172a 0%, #1e293b 60%, #0f172a 100%);
    border: 1px solid #1e293b;
    border-radius: 14px;
    padding: 22px 28px;
    margin-bottom: 18px;
    box-shadow: 0 0 0 1px rgba(56,189,248,0.06), 0 8px 24px rgba(0,0,0,0.35);
}
.hero-title { font-size: 26px; font-weight: 700; color: #f1f5f9; margin: 0; }
.hero-sub { font-size: 13.5px; color: #7dd3fc; margin-top: 4px; font-weight: 500; letter-spacing: 0.02em; text-transform: uppercase; }
.hero-meta { font-size: 13px; color: #94a3b8; margin-top: 10px; }
.hero-meta b { color: #e2e8f0; }

[data-testid="stMetric"] {
    background: #111827;
    border: 1px solid #1f2937;
    border-radius: 10px;
    padding: 12px 14px 8px 14px;
}
[data-testid="stMetricLabel"] { color: #94a3b8 !important; }
[data-testid="stMetricValue"] { color: #f1f5f9 !important; }

section[data-testid="stSidebar"] {
    background: #0b1120;
    border-right: 1px solid #1e293b;
}
div[data-testid="stExpander"] {
    background: #0f172a; border: 1px solid #1e293b !important; border-radius: 10px;
}
hr { border-color: #1e293b; }

/* Tabs */
.stTabs [data-baseweb="tab-list"] { gap: 4px; border-bottom: 1px solid #1e293b; }
.stTabs [data-baseweb="tab"] {
    background: transparent; color: #94a3b8; border-radius: 8px 8px 0 0;
    padding: 8px 16px; font-weight: 600; font-size: 14px;
}
.stTabs [aria-selected="true"] { color: #7dd3fc !important; background: #111827 !important; }

/* Buttons */
.stButton > button, .stDownloadButton > button {
    background: #1e293b; color: #e2e8f0; border: 1px solid #334155;
    border-radius: 8px; font-weight: 600; transition: all 0.15s ease;
}
.stButton > button:hover, .stDownloadButton > button:hover {
    border-color: #38bdf8; color: #7dd3fc; background: #1e293b;
}
.stButton > button[kind="primary"] {
    background: linear-gradient(120deg, #0369a1, #0284c7); border: none; color: #f0f9ff;
}

/* Text / select / multiselect inputs */
.stTextInput input, .stTextArea textarea, .stNumberInput input {
    background: #111827 !important; color: #e2e8f0 !important; border: 1px solid #1f2937 !important; border-radius: 8px !important;
}
.stSelectbox div[data-baseweb="select"] > div, .stMultiSelect div[data-baseweb="select"] > div {
    background: #111827 !important; border: 1px solid #1f2937 !important; border-radius: 8px !important; color: #e2e8f0 !important;
}
[data-baseweb="tag"] { background: #0369a1 !important; }

/* Dataframes */
[data-testid="stDataFrame"] { border: 1px solid #1e293b; border-radius: 10px; overflow: hidden; }

/* File uploader */
[data-testid="stFileUploaderDropzone"] { background: #111827; border: 1px dashed #334155; border-radius: 10px; }

/* Sidebar headers */
section[data-testid="stSidebar"] h1, section[data-testid="stSidebar"] h2, section[data-testid="stSidebar"] h3 {
    color: #f1f5f9 !important;
}

/* Status pill badges used in the checklist */
.status-pill {
    display: inline-block; padding: 3px 11px; border-radius: 999px;
    font-size: 11.5px; font-weight: 700; letter-spacing: 0.02em; text-transform: uppercase;
}
.sev-critical { background: rgba(239,68,68,0.15); color: #fca5a5; border: 1px solid rgba(239,68,68,0.4); }
.sev-high { background: rgba(249,115,22,0.15); color: #fdba74; border: 1px solid rgba(249,115,22,0.4); }
.sev-medium { background: rgba(234,179,8,0.15); color: #fde047; border: 1px solid rgba(234,179,8,0.4); }
.sev-low { background: rgba(148,163,184,0.15); color: #cbd5e1; border: 1px solid rgba(148,163,184,0.4); }

.control-card {
    background: #0d1424; border: 1px solid #1e293b; border-radius: 10px;
    padding: 14px 16px; margin-bottom: 10px;
}
.control-id { color: #38bdf8; font-weight: 700; font-family: 'SFMono-Regular', Consolas, monospace; font-size: 12.5px; }
.control-ref { color: #64748b; font-size: 12px; }
.control-step {
    background: #0b1120; border: 1px solid #1e293b; border-radius: 6px;
    padding: 8px 11px; font-family: 'SFMono-Regular', Consolas, monospace;
    font-size: 12px; color: #a5b4fc; margin: 6px 0;
}
.fw-line { color: #64748b; font-size: 11.5px; }

.library-badge {
    display: inline-flex; align-items: center; gap: 6px;
    background: #111827; border: 1px solid #1f2937; border-radius: 999px;
    padding: 4px 12px; font-size: 12px; color: #94a3b8; margin-bottom: 4px;
}
.status-dot { width: 7px; height: 7px; border-radius: 50%; display: inline-block; }
.dot-on { background: #22c55e; box-shadow: 0 0 6px #22c55e; }
.dot-off { background: #475569; }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# CONTROL CLASSIFICATION - severity + compliance-framework cross-mapping
# Rule-based on category/control keywords rather than hand-tagged per item.
# This is a starting point a reviewer should confirm before client delivery,
# same as the CIS section references elsewhere in this tool - keyword rules
# can misclassify an unusual control, so treat this as triage, not a final
# audit opinion.
# ---------------------------------------------------------------------------

SEVERITY_WEIGHTS = {"Critical": 4, "High": 3, "Medium": 2, "Low": 1}

_SEVERITY_RULES = [
    (4, ["root login", "default password", "empty password", "anonymous", "sa account",
         "enable secret", "encryption key", "unencrypted", "public", "0.0.0.0/0",
         "xp_cmdshell", "anonymous auth", "trust-all", "hardcoded"]),
    (3, ["authentication", "password", "mfa", "multi-factor", "privilege", "admin",
         "encrypt", "tls", "ssl", "firewall", "acl", "access control", "rbac",
         "vault", "secrets", "delegation", "superuser", "sa ", "audit trail"]),
    (2, ["logging", "audit", "monitoring", "patch", "update", "backup", "session",
         "timeout", "ntp", "dns", "certificate", "header"]),
]


def classify_control(category: str, control: str) -> dict:
    text = f"{category} {control}".lower()
    severity = "Low"
    for weight, keywords in _SEVERITY_RULES:
        if any(k in text for k in keywords):
            severity = {4: "Critical", 3: "High", 2: "Medium"}[weight]
            break

    frameworks = {}
    if any(k in text for k in ["auth", "password", "mfa", "admin", "privilege", "access control", "acl", "rbac"]):
        frameworks = {"ISO 27001": "A.9 Access Control", "NIST CSF": "PR.AC", "PCI DSS": "Req 7-8", "SOC 2": "CC6.1"}
    elif any(k in text for k in ["log", "audit", "monitor"]):
        frameworks = {"ISO 27001": "A.12.4 Logging & Monitoring", "NIST CSF": "DE.AE / DE.CM", "PCI DSS": "Req 10", "SOC 2": "CC7.2"}
    elif any(k in text for k in ["encrypt", "tls", "ssl", "certificate"]):
        frameworks = {"ISO 27001": "A.10 Cryptography", "NIST CSF": "PR.DS", "PCI DSS": "Req 3-4", "SOC 2": "CC6.7"}
    elif any(k in text for k in ["firewall", "network", "vpn", "interface", "acl", "flow log"]):
        frameworks = {"ISO 27001": "A.13 Network Security", "NIST CSF": "PR.AC / PR.PT", "PCI DSS": "Req 1", "SOC 2": "CC6.6"}
    elif any(k in text for k in ["patch", "version", "update", "cpu", "installation"]):
        frameworks = {"ISO 27001": "A.12.6 Vulnerability Management", "NIST CSF": "ID.RA / PR.MA", "PCI DSS": "Req 6", "SOC 2": "CC7.1"}
    elif any(k in text for k in ["backup", "recovery", "ha ", "high availability", "replication"]):
        frameworks = {"ISO 27001": "A.17 Business Continuity", "NIST CSF": "PR.IP / RC.RP", "PCI DSS": "Req 12", "SOC 2": "A1.2"}
    else:
        frameworks = {"ISO 27001": "A.12 Operations Security", "NIST CSF": "PR.IP", "PCI DSS": "Req 2", "SOC 2": "CC6.8"}

    return {"severity": severity, "frameworks": frameworks}


def build_docx_report(client_name, technology, reviewer_name, weighted_pct, raw_pct,
                       totals, open_items, exec_summary, all_responses) -> bytes:
    """Build a client-facing DOCX report. Returns raw bytes for st.download_button."""
    if not DOCX_AVAILABLE:
        return b""

    doc = Document()

    title = doc.add_heading("Security Configuration Assessment Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Client: {client_name}\n").bold = True
    meta.add_run(f"Technology Assessed: {technology}\n")
    meta.add_run(f"Reviewer: {reviewer_name or '—'}\n")
    meta.add_run(f"Report Date: {datetime.now().strftime('%Y-%m-%d')}")

    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    if exec_summary:
        doc.add_paragraph(exec_summary)
    else:
        doc.add_paragraph(
            "Executive summary not generated for this report. Use the "
            "'Generate Executive Summary' AI feature on the Compliance "
            "Dashboard tab before exporting to include one here."
        )

    doc.add_heading("Compliance Overview", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Metric", "Value"
    metrics = [
        ("Risk-Weighted Compliance", f"{weighted_pct}%"),
        ("Raw Compliance", f"{raw_pct}%"),
        ("Total Controls Assessed", str(totals["total"])),
        ("Compliant", str(totals["compliant"])),
        ("Non-Compliant", str(totals["noncompliant"])),
        ("Compensating Control", str(totals["compensating"])),
        ("Not Reviewed", str(totals["not_reviewed"])),
    ]
    for label, value in metrics:
        row = table.add_row().cells
        row[0].text, row[1].text = label, value

    doc.add_heading("Open Findings (Remediation Tracker)", level=1)
    if open_items:
        sev_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
        sorted_items = sorted(open_items, key=lambda x: sev_order.get(x.get("Severity"), 4))
        ftable = doc.add_table(rows=1, cols=5)
        ftable.style = "Light Grid Accent 1"
        fhdr = ftable.rows[0].cells
        for idx, label in enumerate(["Item", "Severity", "Control", "Audit Step", "Notes"]):
            fhdr[idx].text = label
        for it in sorted_items:
            row = ftable.add_row().cells
            row[0].text = str(it.get("Item", ""))
            row[1].text = str(it.get("Severity", ""))
            row[2].text = str(it.get("Control", ""))
            row[3].text = str(it.get("Audit Step", ""))
            row[4].text = str(it.get("Notes", "") or "")
    else:
        doc.add_paragraph("No open Non-Compliant findings at time of report generation.")

    doc.add_heading("Full Control Detail", level=1)
    dtable = doc.add_table(rows=1, cols=5)
    dtable.style = "Light Grid Accent 1"
    dhdr = dtable.rows[0].cells
    for idx, label in enumerate(["Item", "Control", "Status", "Severity", "Reference"]):
        dhdr[idx].text = label
    for item_id, data in sorted(all_responses.items()):
        row = dtable.add_row().cells
        row[0].text = item_id
        row[1].text = str(data.get("control", ""))
        row[2].text = str(data.get("status", ""))
        row[3].text = str(data.get("severity", ""))
        row[4].text = str(data.get("reference", ""))

    footer_note = doc.add_paragraph()
    footer_note.add_run(
        "\nThis report reflects a point-in-time assessment. Severity ratings "
        "and framework mappings are generated by rule-based classification "
        "and should be confirmed by the assigned reviewer before client "
        "delivery."
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# 1. CHECKLIST DATA
# Each technology maps to a list of controls: id, category, control text,
# guidance, and a rough CIS Benchmark cross-reference (where a formal CIS
# Benchmark exists for that technology). For technologies without a
# published CIS Benchmark, the list is a generic best-practice checklist
# built from common hardening/compliance frameworks (NIST, ISO 27001,
# vendor hardening guides) rather than a specific named source.
# ---------------------------------------------------------------------------

CHECKLISTS = {
    "Linux Server (CIS Benchmark based)": [
        ("L-1", "Initial Setup & Filesystem", "Separate partitions used for /tmp, /var, /var/log, /var/log/audit, /home", "CIS 1.1.x", "lsblk -f  and  cat /etc/fstab  — confirm /tmp, /var, /var/log, /var/log/audit, /home each have a separate partition entry"),
        ("L-2", "Initial Setup & Filesystem", "Unused filesystems (cramfs, freevxfs, squashfs, udf) are disabled", "CIS 1.1.1.x", "modprobe -n -v cramfs freevxfs squashfs udf  — should show install /bin/true, or check /etc/modprobe.d/ for blacklist entries"),
        ("L-3", "Initial Setup & Filesystem", "GPG keys are configured for package repositories", "CIS 1.2.x", "rpm -q gpg-pubkey --qf '%{name}-%{version}\n' (RHEL) or apt-key list / ls /etc/apt/trusted.gpg.d (Debian)"),
        ("L-4", "Initial Setup & Filesystem", "AppArmor/SELinux is installed, enabled and enforcing", "CIS 1.6.x", "sestatus (SELinux) or aa-status (AppArmor) — confirm status is enforcing/enabled"),
        ("L-5", "Initial Setup & Filesystem", "Bootloader password is set and permissions restricted", "CIS 1.4.x", "cat /boot/grub2/grub.cfg | grep -i password  and  stat -c '%a %U' /boot/grub2/grub.cfg (should be 600, root-owned)"),
        ("L-6", "Services", "Unnecessary services (telnet, rsh, ypserv, tftp, xinetd) are removed/disabled", "CIS 2.1-2.2", "systemctl list-unit-files | grep enabled | grep -E 'telnet|rsh|ypserv|tftp|xinetd'"),
        ("L-7", "Services", "Time synchronization (chrony/ntp) is enabled and configured", "CIS 2.1.1", "systemctl status chronyd  or  timedatectl  — confirm 'NTP synchronized: yes'"),
        ("L-8", "Network Configuration", "IP forwarding is disabled unless the host is a router", "CIS 3.1.x", "sysctl net.ipv4.ip_forward  (expect 0 unless the host is an intended router)"),
        ("L-9", "Network Configuration", "ICMP redirects, source routing are disabled", "CIS 3.2.x", "sysctl net.ipv4.conf.all.accept_redirects net.ipv4.conf.all.send_redirects  (expect 0)"),
        ("L-10", "Network Configuration", "Host-based firewall (nftables/iptables/firewalld) is enabled with default-deny", "CIS 3.5.x", "systemctl status firewalld  or  nft list ruleset  or  iptables -L -n"),
        ("L-11", "Logging & Auditing", "auditd is installed, enabled, and logs are retained per policy", "CIS 4.1.x", "systemctl status auditd  and  auditctl -l"),
        ("L-12", "Logging & Auditing", "rsyslog/journald forwards logs to a central log server", "CIS 4.2.x", "cat /etc/rsyslog.conf /etc/rsyslog.d/*.conf  — check for a remote @@loghost target"),
        ("L-13", "Logging & Auditing", "Logrotate is configured to prevent disk exhaustion", "CIS 4.3", "cat /etc/logrotate.conf /etc/logrotate.d/*"),
        ("L-14", "Access, Authentication & Authorization", "Password complexity, history, and lockout policy enforced via pam_pwquality/faillock", "CIS 5.3.x", "cat /etc/security/pwquality.conf /etc/security/faillock.conf"),
        ("L-15", "Access, Authentication & Authorization", "SSH root login disabled, protocol 2 only, strong ciphers/MACs only", "CIS 5.2.x", "sshd -T | grep -Ei 'permitrootlogin|protocol|ciphers|macs'"),
        ("L-16", "Access, Authentication & Authorization", "sudo requires password and logs all sudo activity", "CIS 5.4.x", "cat /etc/sudoers /etc/sudoers.d/*  and  grep sudo /var/log/secure  (RHEL) or /var/log/auth.log (Debian)"),
        ("L-17", "Access, Authentication & Authorization", "Empty passwords, unused/system accounts are locked or removed", "CIS 5.5.x", "awk -F: '($2 == \"\") {print $1}' /etc/shadow  and  passwd -Sa"),
        ("L-18", "System Maintenance", "World-writable files and unowned files/directories are found and remediated", "CIS 6.1.x", "find / -xdev -type f -perm -0002 2>/dev/null"),
        ("L-19", "System Maintenance", "Permissions on /etc/passwd, /etc/shadow, /etc/gshadow are correctly restricted", "CIS 6.1.x", "stat -c '%a %U:%G %n' /etc/passwd /etc/shadow /etc/gshadow"),
        ("L-20", "System Maintenance", "OS and package patching cadence is documented and current", "CIS 1.9 / general", "yum history  (RHEL) or apt list --upgradable  (Debian), cross-checked against the patch management ticket/SOP"),
    ],
    "Windows Server / Active Directory (CIS Benchmark based)": [
        ("W-1", "Account Policies", "Password policy meets length/complexity/history/max-age requirements", "CIS 1.1.x", "secpol.msc → Account Policies → Password Policy, or  net accounts"),
        ("W-2", "Account Policies", "Account lockout threshold, duration and reset counter are configured", "CIS 1.2.x", "secpol.msc → Account Policies → Account Lockout Policy, or  net accounts"),
        ("W-3", "Local Policies / Audit Policy", "Advanced audit policy is enabled for logon, account management, policy change", "CIS 17.x", "auditpol /get /category:*"),
        ("W-4", "Local Policies / Audit Policy", "User rights assignment restricts 'Log on locally', 'Access this computer from the network' etc.", "CIS 2.2.x", "secpol.msc → Local Policies → User Rights Assignment, or  gpresult /h report.html"),
        ("W-5", "Local Policies / Audit Policy", "Security options: LM hash storage disabled, NTLMv2 only, anonymous enumeration disabled", "CIS 2.3.x", "secpol.msc → Local Policies → Security Options, or  reg query HKLM\\SYSTEM\\CurrentControlSet\\Control\\Lsa"),
        ("W-6", "Windows Firewall", "Domain/Private/Public firewall profiles are all enabled with logging on", "CIS 9.1-9.3", "netsh advfirewall show allprofiles"),
        ("W-7", "Network", "SMBv1 is disabled; SMB signing is required", "CIS 18.3.x", "Get-SmbServerConfiguration | Select EnableSMB1Protocol, RequireSecuritySignature  (PowerShell)"),
        ("W-8", "System Services", "Unnecessary services (Telnet, Remote Registry, SNMP if unused) are disabled", "CIS 5.x", "Get-Service | Where-Object {$_.Status -eq 'Running'}  compared against an approved services baseline, or services.msc"),
        ("W-9", "Administrative Templates", "LAPS or equivalent local admin password rotation is in place", "CIS/MS best practice", "Get-ADComputer -Filter * -Property ms-Mcs-AdmPwdExpirationTime  or confirm LAPS GPO is applied via gpresult"),
        ("W-10", "Administrative Templates", "Windows Update is configured for timely patch deployment", "CIS 18.9.x", "Check WSUS console compliance report, or  Get-WindowsUpdateLog  (PowerShell)"),
        ("W-11", "Administrative Templates", "PowerShell logging (module, script block, transcription) is enabled", "CIS 18.9.100.x", "Get-WinEvent -LogName 'Microsoft-Windows-PowerShell/Operational'  and check ScriptBlockLogging registry key/GPO"),
        ("W-12", "Active Directory", "Privileged groups (Domain Admins, Enterprise Admins) are reviewed and minimized", "AD hardening best practice", "Get-ADGroupMember 'Domain Admins'  and  Get-ADGroupMember 'Enterprise Admins'  (PowerShell/ActiveDirectory module)"),
        ("W-13", "Active Directory", "Kerberos delegation is restricted (no unconstrained delegation on non-critical hosts)", "AD hardening best practice", "Get-ADComputer -Filter {TrustedForDelegation -eq $true}  and  Get-ADUser -Filter {TrustedForDelegation -eq $true}"),
        ("W-14", "Active Directory", "Tiered administration model separates workstation/server/DC admin credentials", "AD hardening best practice", "Review AD OU structure, admin account naming/scoping, and interview the IT team on the documented tiering policy"),
        ("W-15", "Active Directory", "AD backups (including SYSVOL, NTDS.dit) are taken and tested regularly", "General best practice", "wbadmin get versions  and check AD Recycle Bin status; perform/observe a test restore in a lab environment"),
        ("W-16", "Endpoint Protection", "Microsoft Defender / EDR is installed, updated and reporting centrally", "CIS 18.9.x", "Get-MpComputerStatus  (PowerShell) or the EDR console's endpoint status dashboard"),
        ("W-17", "BitLocker / Data Protection", "BitLocker (or equivalent) full-disk encryption is enabled on all endpoints/servers with sensitive data", "CIS 18.9.x", "manage-bde -status  or  Get-BitLockerVolume  (PowerShell)"),
    ],
    "macOS (CIS Benchmark based)": [
        ("MAC-1", "Installation", "Running a currently supported macOS major version with latest security updates applied", "CIS Apple macOS Benchmark 1.1", "sw_vers  and  softwareupdate --history  — compare against Apple's currently supported release list"),
        ("MAC-2", "Filesystem & Encryption", "FileVault full-disk encryption is enabled on all endpoints", "CIS Apple macOS Benchmark 2.5", "fdesetup status  — confirm 'FileVault is On.'"),
        ("MAC-3", "Filesystem & Encryption", "System Integrity Protection (SIP) is enabled", "CIS Apple macOS Benchmark 1.6", "csrutil status  — confirm 'System Integrity Protection status: enabled.'"),
        ("MAC-4", "Network Security", "Application firewall is enabled and stealth mode configured", "CIS Apple macOS Benchmark 2.6", "/usr/libexec/ApplicationFirewall/socketfilterfw --getglobalstate"),
        ("MAC-5", "Network Security", "Remote Login (SSH) and Remote Management are disabled unless explicitly required", "CIS Apple macOS Benchmark 2.4.x", "sudo systemsetup -getremotelogin  and  sudo /System/Library/CoreServices/RemoteManagement/ARDAgent.app/Contents/Resources/kickstart -status"),
        ("MAC-6", "Authentication", "Password policy (length, complexity, expiration) is enforced via MDM configuration profile", "CIS Apple macOS Benchmark 5.x", "pwpolicy -getaccountpolicies  or review the deployed MDM password policy payload"),
        ("MAC-7", "Authentication", "Automatic login is disabled", "CIS Apple macOS Benchmark 5.8", "sudo defaults read /Library/Preferences/com.apple.loginwindow autoLoginUser  (expect no value / error)"),
        ("MAC-8", "Logging & Monitoring", "Unified logging retention meets the organization's minimum window", "CIS Apple macOS Benchmark 3.x", "log config --status  and  log show --last 24h --predicate 'eventType == logEvent'"),
        ("MAC-9", "Software Update", "Automatic checking and installation of security updates is enabled", "CIS Apple macOS Benchmark 1.1", "sudo softwareupdate --schedule  and  defaults read /Library/Preferences/com.apple.SoftwareUpdate"),
        ("MAC-10", "Endpoint Protection", "Managed EDR/antivirus agent is installed, updated, and reporting centrally", "General best practice", "Review the MDM/EDR console's endpoint enrollment and last-check-in status for the device"),
        ("MAC-11", "Gatekeeper", "Gatekeeper is enabled to restrict execution of unsigned/unnotarized applications", "CIS Apple macOS Benchmark 2.6.x", "spctl --status  — confirm 'assessments enabled.'"),
    ],
    "Cloud (AWS / Azure / GCP - CIS Benchmark based)": [
        ("C-1", "Identity & Access Management", "Root/global admin account has MFA enabled and is not used for daily operations", "CIS 1.x", "AWS: IAM console → root user MFA status, or aws iam get-account-summary. Azure: Entra ID → Global Admin MFA. GCP: Cloud Identity admin console"),
        ("C-2", "Identity & Access Management", "IAM policies follow least privilege; no wildcard (*:*) admin policies attached broadly", "CIS 1.x", "AWS: aws iam get-account-authorization-details + IAM Access Analyzer. Azure: Entra ID PIM access reviews. GCP: gcloud projects get-iam-policy"),
        ("C-3", "Identity & Access Management", "Access keys/service credentials are rotated and unused ones are disabled", "CIS 1.x", "AWS: aws iam list-access-keys (check CreateDate). Azure: App registration secrets expiry in Entra ID. GCP: gcloud iam service-accounts keys list"),
        ("C-4", "Identity & Access Management", "Single sign-on / centralized identity provider is used instead of per-service local accounts", "General best practice", "Check the Identity Center (AWS SSO) / Entra ID / Cloud Identity admin console for configured SSO and per-service local account policy"),
        ("C-5", "Logging & Monitoring", "Account-level audit logging (CloudTrail/Activity Log/Cloud Audit Logs) is enabled in all regions", "CIS 2.x/3.x", "AWS: aws cloudtrail describe-trails. Azure: Activity Log settings in the portal. GCP: gcloud logging sinks list"),
        ("C-6", "Logging & Monitoring", "Logs are shipped to a centralized, tamper-evident, access-controlled log store", "CIS 3.x", "Check the CloudTrail S3 bucket / Log Analytics workspace / Cloud Logging export destination and its access-control policy"),
        ("C-7", "Logging & Monitoring", "Alerting is configured for root usage, IAM policy changes, and security group changes", "CIS 4.x", "Check CloudWatch Alarms / Azure Monitor Alert rules / GCP Alerting policies scoped to root usage, IAM, and security-group changes"),
        ("C-8", "Networking", "Default security group / NSG denies all inbound traffic by default", "CIS 5.x", "AWS: aws ec2 describe-security-groups --group-names default. Azure: default NSG rules. GCP: default VPC firewall rules"),
        ("C-9", "Networking", "No security group/NSG rule allows unrestricted inbound access (0.0.0.0/0) on admin ports (22/3389)", "CIS 5.x", "AWS: aws ec2 describe-security-groups filtered for 0.0.0.0/0 on ports 22/3389. Azure/GCP: equivalent NSG/firewall rule review"),
        ("C-10", "Networking", "VPC/VNet flow logs are enabled", "CIS 3.x", "AWS: aws ec2 describe-flow-logs. Azure: Network Watcher → NSG Flow Logs. GCP: gcloud compute networks subnets describe (Flow Logs field)"),
        ("C-11", "Storage & Data Protection", "Object storage buckets/containers are not publicly readable/writable unless explicitly required", "CIS 2.x", "AWS: aws s3api get-bucket-policy-status. Azure: Storage Account → 'Allow Blob public access'. GCP: gsutil iam get gs://<bucket>"),
        ("C-12", "Storage & Data Protection", "Data at rest is encrypted using platform-managed or customer-managed keys", "CIS 2.x", "Check KMS/encryption settings on the storage or database resource in the respective cloud console"),
        ("C-13", "Storage & Data Protection", "Data in transit is enforced over TLS for all managed services", "CIS 2.x", "Check load balancer/API Gateway listener TLS policy configuration in the cloud console"),
        ("C-14", "Monitoring & Response", "A cloud security posture / config-drift tool (Security Hub, Defender for Cloud, Security Command Center) is enabled", "General best practice", "AWS Security Hub console / Microsoft Defender for Cloud / GCP Security Command Center dashboard"),
        ("C-15", "Governance", "Budget/cost anomaly alerts and resource tagging standards are enforced", "General best practice", "AWS Budgets console / Azure Cost Management + Billing / GCP Billing budgets, and a tag-policy compliance report"),
    ],
    "Containers - Docker / Kubernetes (CIS Benchmark based)": [
        ("K-1", "Host Configuration", "Container runtime (Docker/containerd) is kept patched to a supported version", "CIS Docker 1.x", "docker version  /  containerd --version  compared against the vendor's currently supported release list"),
        ("K-2", "Host Configuration", "Docker daemon socket is not exposed over an unauthenticated network port", "CIS Docker 2.x", "netstat -tulpn | grep 2375  (should return nothing) or  docker context ls  to confirm no unauthenticated remote context"),
        ("K-3", "Image & Build", "Images are built from minimal, trusted base images and scanned for CVEs before deployment", "CIS Docker 4.x", "Check the CI/CD pipeline's image-scan step output (Trivy/Grype/Clair) or the registry's vulnerability scan report"),
        ("K-4", "Image & Build", "Containers do not run as root unless explicitly required", "CIS Docker 4.1", "docker inspect <container> --format '{{.Config.User}}'  or Pod spec securityContext.runAsNonRoot"),
        ("K-5", "Image & Build", "Secrets are not baked into images or environment variables in plaintext", "CIS Docker 5.x", "trivy image --scanners secret <image>  or  docker history <image>"),
        ("K-6", "Runtime", "Containers run with a read-only root filesystem where possible", "CIS Docker 5.x", "docker inspect --format '{{.HostConfig.ReadonlyRootfs}}' <container>  or Pod securityContext.readOnlyRootFilesystem"),
        ("K-7", "Runtime", "Resource limits (CPU/memory) are set to prevent noisy-neighbor/DoS", "CIS Docker 5.x", "docker inspect --format '{{.HostConfig.Memory}} {{.HostConfig.NanoCpus}}' <container>  or  kubectl describe pod <pod>  (Resources section)"),
        ("K-8", "Kubernetes - Control Plane", "kube-apiserver anonymous auth is disabled and RBAC authorization mode is enforced", "CIS Kubernetes 1.2.x", "Check kube-apiserver manifest/flags for --anonymous-auth=false and --authorization-mode=RBAC"),
        ("K-9", "Kubernetes - Control Plane", "etcd is encrypted at rest and access is restricted to control-plane nodes only", "CIS Kubernetes 2.x", "Check kube-apiserver's --encryption-provider-config flag, and etcd's network policy/firewall rules restricting access to control-plane nodes"),
        ("K-10", "Kubernetes - Control Plane", "Audit logging is enabled on the API server", "CIS Kubernetes 1.2.x", "Check kube-apiserver's --audit-log-path and --audit-policy-file flags"),
        ("K-11", "Kubernetes - Workloads", "NetworkPolicies restrict pod-to-pod traffic to what's required", "CIS Kubernetes 5.3.x", "kubectl get networkpolicy --all-namespaces"),
        ("K-12", "Kubernetes - Workloads", "PodSecurity admission (or equivalent) blocks privileged/hostPath/hostNetwork pods by default", "CIS Kubernetes 5.2.x", "kubectl get ns -o jsonpath='{.items[*].metadata.labels}'  (pod-security.kubernetes.io labels) or review OPA/Gatekeeper policies"),
        ("K-13", "Kubernetes - Workloads", "Namespaces and RBAC RoleBindings scope access per team/tenant", "CIS Kubernetes 5.1.x", "kubectl get rolebindings,clusterrolebindings --all-namespaces"),
        ("K-14", "Secrets Management", "Kubernetes Secrets are encrypted at rest or an external secrets manager (Vault, KMS) is used", "CIS Kubernetes 2.x / general", "Check the cluster's EncryptionConfiguration resource, or the external secrets manager (Vault/KMS) integration config"),
        ("K-15", "Supply Chain", "Only images from an approved, scanned registry can be deployed (admission control/policy)", "General best practice", "Review the admission controller policy (e.g. Kyverno/OPA image allowlist) or registry webhook configuration"),
    ],
    "VMware vSphere / ESXi (Security Configuration Guide based)": [
        ("VM-1", "Host Access", "ESXi Lockdown Mode is enabled to restrict direct host access to vCenter only", "VMware vSphere Security Configuration Guide", "Host → Configure → Security Profile → Lockdown Mode, or  esxcli system lockdown get"),
        ("VM-2", "Authentication", "vCenter is joined to a centralized identity source (AD/LDAP) rather than relying on local accounts", "VMware vSphere Security Configuration Guide", "vCenter → Administration → Single Sign On → Configuration → Identity Sources"),
        ("VM-3", "Authentication", "ESXi root account use is minimized; named admin accounts with RBAC roles are used instead", "VMware vSphere Security Configuration Guide", "Host → Configure → Users and Groups, or  esxcli system account list"),
        ("VM-4", "Network Security", "Management, vMotion, and storage traffic are isolated on dedicated VLANs/port groups", "VMware vSphere Security Configuration Guide", "Review vSwitch/distributed switch port group VLAN assignments for Management, vMotion, and Storage traffic"),
        ("VM-5", "Network Security", "vMotion traffic is encrypted", "VMware vSphere Security Configuration Guide", "VM → Configure → VMware EVC/Edit Settings → check 'Encrypted vMotion' setting"),
        ("VM-6", "Logging & Monitoring", "ESXi hosts and vCenter forward logs to a central syslog server", "VMware vSphere Security Configuration Guide", "Host → Configure → Advanced System Settings → Syslog.global.logHost, or  esxcli system syslog config get"),
        ("VM-7", "Patch Management", "Hosts, vCenter, and VMware Tools are on a currently supported, patched version", "VMware vSphere Security Configuration Guide", "vCenter → Updates tab, or  esxcli system version get  compared against VMware's current release/EOL list"),
        ("VM-8", "Certificates", "vCenter and ESXi hosts use certificates issued by a trusted internal or external CA, not defaults", "VMware vSphere Security Configuration Guide", "vCenter → Administration → Certificate Management, or  esxcli system security certificatestore list"),
        ("VM-9", "Snapshots & Backup", "VM backups are encrypted, tested via restore, and snapshot sprawl is monitored", "General best practice", "Review the backup solution's job configuration/encryption settings and the most recent test-restore log; check for long-lived snapshots"),
        ("VM-10", "Access Control", "Role-based access control assigns least-privilege permissions per team/tenant at the folder/resource-pool level", "VMware vSphere Security Configuration Guide", "vCenter → Administration → Access Control → Roles, and review permission assignments per object"),
    ],
    "Microsoft 365 (Security Baseline based)": [
        ("M365-1", "Identity", "Conditional Access or security defaults enforce MFA for all users", "Microsoft 365 Security Baseline", "Entra admin center → Protection → Conditional Access → Policies, or Microsoft Entra security defaults status"),
        ("M365-2", "Identity", "Legacy authentication protocols (POP, IMAP, basic auth) are blocked", "Microsoft 365 Security Baseline", "Entra admin center → Conditional Access policy blocking legacy authentication, or  Get-CASMailboxPlan  (Exchange Online PowerShell)"),
        ("M365-3", "Email Security", "DMARC, DKIM, and SPF are enforced for all organizational sending domains", "Microsoft 365 Security Baseline", "nslookup -type=TXT _dmarc.<domain>  and check DKIM/SPF records via DNS lookup"),
        ("M365-4", "Email Security", "Safe Links and Safe Attachments (Defender for Office 365) are enabled", "Microsoft 365 Security Baseline", "Microsoft Defender portal → Email & collaboration → Policies & rules → Threat policies"),
        ("M365-5", "Data Protection", "Sensitivity labels and DLP policies are applied to sensitive content", "Microsoft Purview Security Baseline", "Purview compliance portal → Data loss prevention → Policies, and Information Protection → Labels"),
        ("M365-6", "Access Control", "The Global Administrator role is limited to a minimal, monitored set of accounts, with PIM for just-in-time elevation", "Microsoft 365 Security Baseline", "Entra admin center → Roles and administrators → Global Administrator, and Privileged Identity Management assignments"),
        ("M365-7", "Device Management", "Intune compliance policies require managed, encrypted, up-to-date devices before granting access", "Microsoft 365 Security Baseline", "Intune admin center → Devices → Compliance policies, cross-referenced with the Conditional Access grant controls"),
        ("M365-8", "Logging & Monitoring", "The Unified Audit Log is enabled with a defined retention period", "Microsoft 365 Security Baseline", "Purview compliance portal → Audit → Search, or  Get-AdminAuditLogConfig  (Exchange Online PowerShell)"),
        ("M365-9", "External Sharing", "External sharing for SharePoint/OneDrive is restricted to specific domains or disabled by default", "Microsoft 365 Security Baseline", "SharePoint admin center → Policies → Sharing"),
        ("M365-10", "Application Access", "Third-party OAuth app consent is restricted or requires admin approval", "Microsoft 365 Security Baseline", "Entra admin center → Enterprise applications → Consent and permissions → User consent settings"),
    ],
    "Google Workspace (Security Checklist based)": [
        ("GWS-1", "Identity", "2-Step Verification is enforced organization-wide", "Google Workspace Security Checklist", "Admin console → Security → Authentication → 2-step verification"),
        ("GWS-2", "Identity", "Super Admin accounts require security-key-based 2SV and are limited in number", "Google Workspace Security Checklist", "Admin console → Account → Admin roles → Super Admin, and 2SV enrollment status per admin"),
        ("GWS-3", "Email Security", "DMARC, DKIM, and SPF are configured for all sending domains", "Google Workspace Security Checklist", "nslookup -type=TXT _dmarc.<domain>  and check DKIM/SPF records via DNS lookup"),
        ("GWS-4", "Data Sharing", "External file sharing defaults are restricted rather than 'anyone with the link'", "Google Workspace Security Checklist", "Admin console → Apps → Google Workspace → Drive and Docs → Sharing settings"),
        ("GWS-5", "Data Protection", "Data Loss Prevention (DLP) rules are configured for Drive and Gmail", "Google Workspace Security Checklist", "Admin console → Security → Data protection → DLP rules"),
        ("GWS-6", "Device Management", "Mobile/endpoint management requires screen lock and encryption on devices accessing company data", "Google Workspace Security Checklist", "Admin console → Devices → Mobile & endpoints → Settings → Universal settings"),
        ("GWS-7", "Logging & Monitoring", "Admin console audit logs are exported to a SIEM", "Google Workspace Security Checklist", "Admin console → Reporting → Audit and investigation, and the configured BigQuery/Reports API export"),
        ("GWS-8", "Application Access", "Third-party app access (OAuth) is reviewed and restricted to allow-listed apps where appropriate", "Google Workspace Security Checklist", "Admin console → Security → API controls → App access control"),
        ("GWS-9", "Session Security", "Session length and re-authentication policies are configured for sensitive apps", "Google Workspace Security Checklist", "Admin console → Security → Google session control / Context-Aware Access"),
    ],
    "Web Application / API (generic - no single CIS benchmark)": [
        ("A-1", "Authentication & Session", "MFA is available/enforced for privileged and customer-facing accounts", "OWASP ASVS aligned", "Attempt a login flow manually and review the IAM/Auth0/Okta MFA enforcement policy"),
        ("A-2", "Authentication & Session", "Session tokens are random, short-lived, and invalidated on logout/password change", "OWASP ASVS aligned", "Inspect session cookie flags (HttpOnly, Secure, SameSite) via browser dev tools or Burp Suite; check token expiry in code/config"),
        ("A-3", "Input Handling", "All user input is validated/sanitized server-side; parameterized queries used (no raw SQL concatenation)", "OWASP Top 10 aligned", "Manual code review of the data-access layer for parameterized queries, plus a SAST/DAST scan (e.g. OWASP ZAP, Burp) for injection"),
        ("A-4", "Access Control", "Authorization checks are enforced server-side on every request, not just in the UI", "OWASP Top 10 aligned", "Manual pentest: attempt to access another user's resource by changing an ID/parameter (IDOR/privilege-escalation test)"),
        ("A-5", "Transport & Headers", "TLS 1.2+ enforced; HSTS, CSP, X-Content-Type-Options headers set", "OWASP Secure Headers", "curl -I https://target  and an SSL Labs (ssllabs.com) or securityheaders.com scan"),
        ("A-6", "Secrets & Config", "API keys/secrets are stored in a vault/secret manager, not in source control or client-side code", "General best practice", "gitleaks detect  or  trufflehog  against the source repo; confirm vault/secret-manager integration in config"),
        ("A-7", "Logging & Monitoring", "Security-relevant events (auth failures, privilege changes) are logged without leaking sensitive data", "OWASP logging cheat sheet aligned", "Review log output for auth-failure/privilege-change events and confirm no PII/secrets appear in log samples"),
        ("A-8", "Dependency Management", "Third-party libraries/dependencies are scanned for known vulnerabilities on a regular cadence", "General best practice", "Check the CI pipeline for npm audit / pip-audit / Snyk / Dependabot output"),
        ("A-9", "Rate Limiting & Abuse", "Rate limiting / throttling is applied to authentication and other sensitive endpoints", "General best practice", "Send a burst of requests to the login endpoint and observe the throttling response, or review the API gateway's rate-limit config"),
        ("A-10", "Data Protection", "PII/sensitive fields are encrypted at rest and masked in logs/non-prod environments", "General best practice", "Review the DB schema/column-level encryption config and check log samples for masked sensitive fields"),
    ],
    "MySQL / Oracle MySQL (CIS Benchmark based)": [
        ("MY-1", "Installation & Patching", "Running a currently supported MySQL version; unused example/test databases removed", "CIS 1.x", "mysql -e 'SELECT VERSION(); SHOW DATABASES;'  (check for a leftover 'test' database)"),
        ("MY-2", "Installation & Patching", "MySQL service runs under a dedicated non-root OS account", "CIS 1.x", "ps -ef | grep mysqld  — confirm the running OS user is not root"),
        ("MY-3", "File Permissions", "Data directory and my.cnf permissions restricted to the mysql OS user only", "CIS 2.x", "ls -la /var/lib/mysql /etc/my.cnf"),
        ("MY-4", "File Permissions", "Error/general logs are not world-readable", "CIS 2.x", "mysql -e 'SHOW VARIABLES LIKE 'log_error';'  then  ls -la <path returned>"),
        ("MY-5", "General & Network", "TLS is enabled and required for client connections (require_secure_transport)", "CIS 3.x", "mysql -e 'SHOW VARIABLES LIKE 'require_secure_transport';'"),
        ("MY-6", "General & Network", "local_infile is disabled unless explicitly required", "CIS 3.x", "mysql -e 'SHOW VARIABLES LIKE 'local_infile';'"),
        ("MY-7", "Authentication", "Default/anonymous accounts are removed or locked; root account has a strong password", "CIS 4.x", "mysql -e 'SELECT User,Host FROM mysql.user;'"),
        ("MY-8", "Authentication", "Password validation plugin enforces complexity, history and expiration", "CIS 4.x", "mysql -e 'SHOW VARIABLES LIKE 'validate_password%';'"),
        ("MY-9", "Authorization", "SUPER, FILE, and PROCESS privileges are limited to admin/service accounts only", "CIS 4.x", "mysql -e 'SELECT User,Host,Super_priv,File_priv FROM mysql.user;'"),
        ("MY-10", "Auditing", "Audit plugin or general query log captures security-relevant events", "CIS 6.x", "mysql -e 'SHOW PLUGINS;'  (check audit_log)  or  'SHOW VARIABLES LIKE 'general_log%';'"),
        ("MY-11", "Auditing", "Logs are forwarded to a central log server/SIEM", "CIS 6.x", "Check the rsyslog/Filebeat configuration shipping MySQL logs to the central SIEM"),
        ("MY-12", "Replication", "Replication traffic between source and replicas is encrypted, if replication is used", "CIS 7.x", "mysql -e 'SHOW SLAVE STATUS\\G'  (check Master_SSL_Allowed)"),
    ],
    "PostgreSQL (CIS Benchmark based)": [
        ("PG-1", "Installation", "Running a currently supported PostgreSQL major version", "CIS 1.x", "psql -c 'SELECT version();'"),
        ("PG-2", "Directory & File Permissions", "PGDATA directory is owned by the postgres user with mode 0700", "CIS 2.x", "ls -ld $PGDATA  (expect drwx------ postgres:postgres)"),
        ("PG-3", "Directory & File Permissions", "pg_hba.conf restricts connections by host/user/database (no trust-all rules)", "CIS 2.x / 6.x", "cat $PGDATA/pg_hba.conf"),
        ("PG-4", "Logging & Auditing", "log_connections, log_disconnections and log_statement are configured per policy", "CIS 3.x", "psql -c 'SHOW log_connections; SHOW log_disconnections; SHOW log_statement;'"),
        ("PG-5", "Logging & Auditing", "log_line_prefix includes timestamp, user, database and session information", "CIS 3.x", "psql -c 'SHOW log_line_prefix;'"),
        ("PG-6", "User Access & Authorization", "Superuser role membership is restricted to DBAs only", "CIS 4.x", "psql -c 'SELECT usename FROM pg_user WHERE usesuper = true;'"),
        ("PG-7", "User Access & Authorization", "Password encryption method is set to scram-sha-256", "CIS 4.x", "psql -c 'SHOW password_encryption;'"),
        ("PG-8", "Connection Security", "SSL is enabled (ssl=on) with valid certificates", "CIS 5.x", "psql -c 'SHOW ssl;'  and  openssl x509 -in server.crt -noout -dates"),
        ("PG-9", "Connection Security", "Idle session / connection timeout is configured", "CIS 5.x", "psql -c 'SHOW idle_in_transaction_session_timeout;'"),
        ("PG-10", "Replication & Backup", "WAL archiving/backups are encrypted at rest and restore has been tested", "CIS 7.x", "Review pg_basebackup/WAL-E/WAL-G/pgBackRest config for encryption flags; perform a test restore"),
    ],
    "Microsoft SQL Server (CIS Benchmark based)": [
        ("MS-1", "Installation & Patching", "SQL Server is on a supported version with the latest cumulative update applied", "CIS 1.x", "SELECT @@VERSION;  in SSMS or sqlcmd"),
        ("MS-2", "Surface Area Reduction", "xp_cmdshell is disabled unless explicitly required and documented", "CIS 2.x", "EXEC sp_configure 'xp_cmdshell';"),
        ("MS-3", "Surface Area Reduction", "Ad Hoc Distributed Queries and OLE Automation Procedures are disabled", "CIS 2.x", "EXEC sp_configure 'Ad Hoc Distributed Queries'; EXEC sp_configure 'Ole Automation Procedures';"),
        ("MS-4", "Authentication", "Windows Authentication is used where feasible; Mixed Mode is justified if enabled", "CIS 3.x", "SELECT SERVERPROPERTY('IsIntegratedSecurityOnly');  or SSMS → Server Properties → Security"),
        ("MS-5", "Authentication", "The sa account is disabled or renamed and has a strong, unique password", "CIS 3.x", "SELECT name, is_disabled FROM sys.sql_logins WHERE name = 'sa';"),
        ("MS-6", "Authorization", "Server and database roles follow least privilege; public role permissions are minimized", "CIS 4.x", "SELECT * FROM sys.database_role_members;  and review permissions granted to the public role"),
        ("MS-7", "Auditing", "SQL Server Audit (or equivalent) is enabled for logins, schema and permission changes", "CIS 5.x", "SELECT * FROM sys.server_audits;  and  SELECT * FROM sys.dm_server_audit_status;"),
        ("MS-8", "Encryption", "Transparent Data Encryption (TDE) is enabled for databases holding sensitive data", "CIS 6.x", "SELECT name, is_encrypted FROM sys.databases;"),
        ("MS-9", "Encryption", "Force Encryption is enabled at the instance level for client connections", "CIS 6.x", "SQL Server Configuration Manager → Protocols → Force Encryption, or SELECT encrypt_option FROM sys.dm_exec_connections;"),
        ("MS-10", "Backup & Recovery", "Backups are encrypted and a restore has been tested within the last review cycle", "General best practice", "RESTORE HEADERONLY FROM DISK = '<backup file>'  (check KeyAlgorithm column) and review backup job logs"),
    ],
    "Oracle Database (CIS Benchmark based)": [
        ("OR-1", "Installation & Patching", "Latest Oracle Critical Patch Update (CPU) is applied", "CIS 1.x", "SELECT * FROM dba_registry_history;  or  opatch lsinventory"),
        ("OR-2", "Listener Security", "Listener has a password/ADMIN_RESTRICTIONS set; remote admin of the listener is disabled", "CIS 2.x", "lsnrctl status  and check listener.ora for PASSWORD/ADMIN_RESTRICTIONS entries"),
        ("OR-3", "Authentication", "Default demo accounts (e.g. SCOTT, DBSNMP) are locked or removed if unused", "CIS 3.x", "SELECT username, account_status FROM dba_users WHERE username IN ('SCOTT','DBSNMP');"),
        ("OR-4", "Authentication", "Password profile enforces complexity, expiration and failed-login lockout", "CIS 3.x", "SELECT * FROM dba_profiles WHERE profile = 'DEFAULT' AND resource_type = 'PASSWORD';"),
        ("OR-5", "Authorization", "PUBLIC role grants are minimized (no broad EXECUTE ANY / SELECT ANY)", "CIS 4.x", "SELECT * FROM dba_tab_privs WHERE grantee = 'PUBLIC' AND privilege LIKE '%ANY%';"),
        ("OR-6", "Auditing", "Unified Audit / fine-grained auditing is enabled for privileged actions", "CIS 5.x", "SELECT * FROM audit_unified_enabled_policies;"),
        ("OR-7", "Auditing", "Audit trail is stored securely and forwarded to a central log server/SIEM", "CIS 5.x", "SELECT * FROM dba_audit_trail WHERE ROWNUM <= 10;  and check forwarding/export configuration"),
        ("OR-8", "Encryption", "Transparent Data Encryption (TDE) is enabled for tablespaces with sensitive data", "CIS 6.x", "SELECT * FROM v$encryption_wallet;  and  SELECT tablespace_name, encrypted FROM dba_tablespaces;"),
        ("OR-9", "Network", "SQL*Net native encryption or TLS is enforced for client-server traffic", "CIS 6.x", "SELECT network_service_banner FROM v$session_connect_info WHERE ROWNUM <=5;  or check sqlnet.ora"),
        ("OR-10", "Backup & Recovery", "RMAN backups are encrypted with a documented, tested retention policy", "General best practice", "RMAN> LIST BACKUP SUMMARY;  (check ENCRYPTION column) and review the most recent restore-test log"),
    ],
    "Apache HTTP Server (CIS Benchmark based)": [
        ("AP-1", "Installation & Modules", "Running a supported Apache version; mod_status/mod_info disabled or access-restricted", "CIS 1.x / 2.x", "apachectl -v  and  apachectl -M | grep -E 'status_module|info_module'"),
        ("AP-2", "Modules", "Unnecessary modules (autoindex, userdir, unused cgi) are disabled", "CIS 2.x", "apachectl -M | grep -E 'autoindex|userdir|cgi'"),
        ("AP-3", "Access Control", "Directory listing is disabled globally (Options -Indexes)", "CIS 3.x", "grep -r 'Options' /etc/httpd/conf* /etc/apache2/  — confirm -Indexes is set"),
        ("AP-4", "Information Disclosure", "ServerTokens is set to Prod and ServerSignature is off", "CIS 3.x", "grep -i servertokens /etc/httpd/conf/httpd.conf  and  curl -I http://target  (check Server header)"),
        ("AP-5", "SSL/TLS", "TLS 1.2+ only, weak ciphers disabled, HSTS header set", "CIS 4.x", "apachectl -M | grep ssl  and  openssl s_client -connect host:443  or an SSL Labs scan"),
        ("AP-6", "Logging", "Access and error logs are enabled with sufficient detail and retention", "CIS 5.x", "grep -i logformat /etc/httpd/conf/httpd.conf"),
        ("AP-7", "Logging", "Logs are forwarded to a central log server/SIEM", "CIS 5.x", "Check the rsyslog/Filebeat configuration shipping Apache logs to the central SIEM"),
        ("AP-8", "Process Security", "Apache runs as a dedicated non-root user (e.g. www-data)", "CIS 6.x", "ps aux | grep httpd  and check the 'User' directive in httpd.conf"),
        ("AP-9", "Request Limits", "Timeout, LimitRequestBody and related directives are configured to mitigate DoS", "CIS 6.x", "grep -iE 'Timeout|LimitRequestBody' /etc/httpd/conf/httpd.conf"),
        ("AP-10", "File Permissions", "Web root and configuration file permissions are restricted to the web server user/admins", "CIS 6.x", "ls -la /etc/httpd/conf/httpd.conf /var/www/html"),
    ],
    "NGINX (CIS Benchmark based)": [
        ("NG-1", "Installation", "Running a supported NGINX version; server_tokens is off", "CIS 1.x / 2.x", "nginx -v  and  grep server_tokens /etc/nginx/nginx.conf"),
        ("NG-2", "Access Control", "autoindex is off globally unless explicitly required", "CIS 3.x", "grep -r autoindex /etc/nginx/"),
        ("NG-3", "SSL/TLS", "TLS 1.2/1.3 only, strong cipher suite, HSTS enabled", "CIS 4.x", "nginx -T | grep ssl_protocols  and an SSL Labs scan"),
        ("NG-4", "Access Control", "client_max_body_size and similar limits are set to prevent abuse", "CIS 3.x", "grep client_max_body_size /etc/nginx/nginx.conf"),
        ("NG-5", "Logging", "access_log and error_log are enabled with adequate detail", "CIS 5.x", "grep -E 'access_log|error_log' /etc/nginx/nginx.conf"),
        ("NG-6", "Logging", "Logs are shipped to a centralized log server/SIEM", "CIS 5.x", "Check the log-shipping agent (Filebeat/rsyslog) configuration"),
        ("NG-7", "Process Security", "Worker processes run as a dedicated non-root user", "CIS 6.x", "ps aux | grep nginx  and check the 'user' directive in nginx.conf"),
        ("NG-8", "Rate Limiting", "limit_req/limit_conn are configured on sensitive endpoints", "CIS 6.x", "grep -rE 'limit_req|limit_conn' /etc/nginx/"),
        ("NG-9", "Security Headers", "X-Content-Type-Options, X-Frame-Options and CSP headers are set", "CIS 4.x / general", "curl -I https://target  — check response headers"),
        ("NG-10", "File Permissions", "Config files and web root are restricted to the nginx user/admins only", "CIS 6.x", "ls -la /etc/nginx/nginx.conf /usr/share/nginx/html"),
    ],
    "Microsoft IIS (CIS Benchmark based)": [
        ("IIS-1", "Installation", "Running a supported IIS/Windows Server version; unused server roles removed", "CIS 1.x", "Get-WindowsFeature Web-Server  (PowerShell) or IIS Manager → Server info"),
        ("IIS-2", "Access Control", "Directory browsing is disabled site-wide", "CIS 3.x", "Get-WebConfigurationProperty -Filter /system.webServer/directoryBrowse -Name enabled  or IIS Manager → Directory Browsing feature"),
        ("IIS-3", "Request Filtering", "Request Filtering module is configured (file extensions, verbs, headers)", "CIS 4.x", "IIS Manager → Request Filtering, or  Get-WebConfiguration //requestFiltering"),
        ("IIS-4", "SSL/TLS", "TLS 1.2+ enforced; weak cipher suites disabled", "CIS 5.x", "Get-TlsCipherSuite  (PowerShell), IIS Crypto tool output, or an SSL Labs scan"),
        ("IIS-5", "Logging", "W3C logging is enabled with adequate fields (client IP, user agent, status)", "CIS 6.x", "IIS Manager → Logging feature, or  Get-WebConfigurationProperty -Filter /system.applicationHost/sites/siteDefaults/logFile"),
        ("IIS-6", "Logging", "Logs are forwarded to a central log server/SIEM", "CIS 6.x", "Check the log-shipping agent (Filebeat/NXLog) configuration"),
        ("IIS-7", "Identity", "Application pools run under least-privilege service accounts, not LocalSystem", "CIS 7.x", "Get-IISAppPool | Select Name, ProcessModel  (PowerShell)"),
        ("IIS-8", "Information Disclosure", "Server header/version info is suppressed", "CIS 3.x", "curl -I https://target  — check Server/X-Powered-By headers"),
        ("IIS-9", "Authentication", "Anonymous authentication is disabled where not required; strong auth enforced", "CIS 7.x", "IIS Manager → Authentication feature, or  Get-WebConfigurationProperty -Filter /system.webServer/security/authentication/anonymousAuthentication -Name enabled"),
        ("IIS-10", "Configuration Security", "machineKey and web.config secrets are encrypted/protected", "CIS 8.x", "aspnet_regiis -pef  output, or check web.config for encrypted <connectionStrings>/<appSettings> sections"),
    ],
    "Apache Tomcat (CIS Benchmark based)": [
        ("TC-1", "Installation", "Running a supported Tomcat version; sample apps (examples, docs, host-manager) removed", "CIS 1.x / 2.x", "catalina.sh version  and  ls $CATALINA_HOME/webapps  (check for examples, docs, host-manager)"),
        ("TC-2", "Access Control", "Manager/Host Manager app access is restricted by IP with strong credentials", "CIS 3.x", "cat $CATALINA_HOME/webapps/manager/META-INF/context.xml  — check for a RemoteAddrValve restriction"),
        ("TC-3", "Authentication", "tomcat-users.xml uses strong, unique passwords; unused roles removed", "CIS 3.x", "cat $CATALINA_HOME/conf/tomcat-users.xml"),
        ("TC-4", "SSL/TLS", "TLS 1.2+ connector configured; weak ciphers disabled", "CIS 4.x", "grep -A5 'Connector.*8443' $CATALINA_HOME/conf/server.xml"),
        ("TC-5", "Logging", "AccessLogValve is enabled with sufficient detail", "CIS 5.x", "grep -A3 AccessLogValve $CATALINA_HOME/conf/server.xml"),
        ("TC-6", "Logging", "Catalina/access logs are forwarded to a central log server/SIEM", "CIS 5.x", "Check the log-shipping agent configuration for catalina.out/localhost_access_log"),
        ("TC-7", "Process Security", "Tomcat runs as a dedicated non-root/non-admin service account", "CIS 6.x", "ps aux | grep catalina  and check the systemd service's User= directive"),
        ("TC-8", "Shutdown Port", "Shutdown port is disabled or bound to localhost with a strong shutdown command", "CIS 2.x", "grep 'Server port' $CATALINA_HOME/conf/server.xml"),
        ("TC-9", "Security Headers", "Server header suppressed; security headers set via filter", "CIS 4.x", "curl -I https://target:8443  and check the security-header filter config in web.xml"),
        ("TC-10", "File Permissions", "CATALINA_HOME/webapps permissions restricted to service account/admins", "CIS 6.x", "ls -la $CATALINA_HOME/webapps $CATALINA_HOME/conf"),
    ],
    "FortiGate Firewall (CIS Benchmark based)": [
        ("FG-1", "Administrative Access", "HTTPS-only admin access; HTTP admin access disabled or redirected", "CIS FortiGate Benchmark", "CLI: show system interface  (check allowaccess field). GUI: System > Interfaces > edit interface"),
        ("FG-2", "Administrative Access", "Admin access restricted to trusted management subnets/interfaces (trusted hosts)", "CIS FortiGate Benchmark", "CLI: show system admin  (check trusthost1-10). GUI: System > Administrators > edit admin"),
        ("FG-3", "Administrative Access", "Idle admin session timeout configured (5 minutes or less recommended)", "CIS FortiGate Benchmark", "CLI: show system global | grep admintimeout. GUI: System > Settings > Administration Settings"),
        ("FG-4", "Authentication", "Local admin accounts use a strong password policy; central auth (LDAP/RADIUS) + MFA enabled", "CIS FortiGate Benchmark", "CLI: show system admin  and  show user radius / show user ldap. GUI: System > Administrators"),
        ("FG-5", "System", "NTP is configured with trusted, authenticated time sources", "CIS FortiGate Benchmark", "CLI: show system ntp. GUI: System > Settings > System Time"),
        ("FG-6", "System", "DNS servers are set to trusted internal/organizational resolvers", "CIS FortiGate Benchmark", "CLI: show system dns. GUI: Network > DNS"),
        ("FG-7", "Logging", "Local and remote logging (FortiAnalyzer/syslog) enabled for traffic, event and admin logs", "CIS FortiGate Benchmark", "CLI: show log syslogd setting  and  show log fortianalyzer setting. GUI: Log & Report > Log Settings"),
        ("FG-8", "Logging", "Log retention meets policy/regulatory requirement", "CIS FortiGate Benchmark", "GUI: Log & Report > Log Settings, or the FortiAnalyzer retention policy console"),
        ("FG-9", "Firewall Policy", "Default deny-all policy exists at the bottom of the policy list; broad/unused policies reviewed", "CIS FortiGate Benchmark", "CLI: show firewall policy. GUI: Policy & Objects > Firewall Policy (check the bottom-most entry)"),
        ("FG-10", "Firewall Policy", "All policies log traffic (log-traffic enabled) for allowed and denied sessions", "CIS FortiGate Benchmark", "CLI: show firewall policy  (check the 'logtraffic' field per policy)"),
        ("FG-11", "VPN", "IPsec/SSL VPN uses strong encryption (AES-256, SHA-256+) with certificate-based or MFA auth", "CIS FortiGate Benchmark", "CLI: show vpn ipsec phase1-interface  and  show vpn ssl settings. GUI: VPN > IPsec/SSL-VPN Settings"),
        ("FG-12", "HA & Certificates", "HA heartbeat interfaces isolated; management certificates are valid and not default self-signed", "CIS FortiGate Benchmark / general", "CLI: show system ha  and  show vpn certificate local. GUI: System > HA, System > Certificates"),
    ],
    "Palo Alto Networks Firewall (CIS Benchmark based)": [
        ("PA-1", "Administrative Access", "Management interface access restricted to specific IPs, HTTPS only", "CIS Palo Alto Benchmark", "CLI: show deviceconfig system | match permitted-ip. GUI: Device > Setup > Management > Interfaces"),
        ("PA-2", "Administrative Access", "Idle timeout for admin sessions is configured", "CIS Palo Alto Benchmark", "CLI: show deviceconfig system | match idle-timeout. GUI: Device > Setup > Management"),
        ("PA-3", "Authentication", "Local admin passwords meet complexity policy; MFA/RADIUS-TACACS+ used for admin auth", "CIS Palo Alto Benchmark", "CLI: show admins. GUI: Device > Administrators, Device > Authentication Profile"),
        ("PA-4", "System", "NTP and DNS configured with trusted sources", "CIS Palo Alto Benchmark", "CLI: show ntp / show dns. GUI: Device > Setup > Services"),
        ("PA-5", "Logging", "Log forwarding to Panorama/syslog configured for traffic, threat, and system logs", "CIS Palo Alto Benchmark", "CLI: show logging-status. GUI: Objects > Log Forwarding, Device > Log Settings"),
        ("PA-6", "Logging", "Log storage/retention meets policy", "CIS Palo Alto Benchmark", "GUI: Device > Setup > Management > Logging and Reporting Settings"),
        ("PA-7", "Security Policy", "Default deny rule present; broad any/any rules reviewed and minimized", "CIS Palo Alto Benchmark", "CLI: show running security-policy. GUI: Policies > Security (review the full rule base top to bottom)"),
        ("PA-8", "Security Policy", "Security profiles (AV, anti-spyware, vulnerability protection) attached to allow rules", "CIS Palo Alto Benchmark", "GUI: Policies > Security — check the Profile column for each rule"),
        ("PA-9", "Certificates", "Management/SSL-decryption certificates are valid and from a trusted CA where applicable", "CIS Palo Alto Benchmark", "CLI: show system state filter-pretty cfg.filter. GUI: Device > Certificate Management > Certificates"),
        ("PA-10", "High Availability", "HA configuration reviewed for failover and sync integrity, if applicable", "CIS Palo Alto Benchmark / general", "CLI: show high-availability state. GUI: Device > High Availability"),
    ],
    "Cisco ASA / IOS (CIS Benchmark based)": [
        ("CS-1", "Administrative Access", "SSH only for management (Telnet disabled); ACL restricts management access to trusted hosts", "CIS Cisco Benchmark", "show running-config | include line vty|ssh|access-class"),
        ("CS-2", "Authentication", "AAA (TACACS+/RADIUS) configured for admin authentication and command authorization", "CIS Cisco Benchmark", "show running-config | include aaa"),
        ("CS-3", "Authentication", "Enable secret (not enable password) is used with a strong, unique value", "CIS Cisco Benchmark", "show running-config | include enable secret"),
        ("CS-4", "Password Policy", "Minimum password length enforced; service password-encryption enabled", "CIS Cisco Benchmark", "show running-config all | include password"),
        ("CS-5", "Logging", "Logging enabled to a central syslog server with timestamps and appropriate severity", "CIS Cisco Benchmark", "show logging  and  show running-config | include logging"),
        ("CS-6", "Logging", "NTP configured and authenticated for accurate log timestamps", "CIS Cisco Benchmark", "show ntp status  and  show running-config | include ntp"),
        ("CS-7", "Interfaces", "Unused interfaces are administratively shut down", "CIS Cisco Benchmark", "show ip interface brief  (check for administratively down status)"),
        ("CS-8", "Access Control", "ACLs follow least privilege with explicit deny-log rules where needed", "CIS Cisco Benchmark", "show access-lists"),
        ("CS-9", "Services", "Unnecessary services (CDP, HTTP server, small services) disabled on untrusted-facing interfaces", "CIS Cisco Benchmark", "show running-config | include cdp run|ip http server"),
        ("CS-10", "SNMP", "SNMPv1/v2c disabled or restricted; SNMPv3 with authPriv used if SNMP is required", "CIS Cisco Benchmark", "show running-config | include snmp-server"),
        ("CS-11", "Banner", "Legal/warning login banner is configured", "CIS Cisco Benchmark", "show running-config | include banner"),
    ],
    "IDS/IPS - Snort / Suricata (generic - no single CIS benchmark)": [
        ("ID-1", "Deployment", "Sensor placed at appropriate network choke point(s), covering all critical traffic paths", "NIST SP 800-94", "Review the network diagram plus SPAN/TAP or inline bridge configuration (ip link show for bridge interfaces)"),
        ("ID-2", "Rule Management", "Rule sets (ET Open/Pro, Talos, custom) updated on a defined, automated schedule", "Vendor best practice", "Check suricata-update / PulledPork last-run timestamp, or the cron job scheduling rule updates"),
        ("ID-3", "Rule Management", "Custom rules reviewed for false-positive rate before enabling in blocking mode", "Vendor best practice", "Review local.rules and the alert log's false-positive rate over the trial/tuning period"),
        ("ID-4", "Performance", "Sensor sized/tuned to avoid packet drops under peak traffic load (verified via stats/perf counters)", "Vendor best practice", "suricatasc -c stats  or the Snort perfmon output; check dropped-packet counters"),
        ("ID-5", "Logging & Alerting", "Alerts forwarded to a central SIEM with sufficient context; pcap retention policy defined", "NIST SP 800-94", "Check the eve.json/unified2 output configuration and the SIEM ingestion pipeline status"),
        ("ID-6", "Operating Mode", "IPS/blocking mode enabled only after a tuning period in IDS/alert-only mode", "Vendor best practice", "Check the af-packet/nfq mode setting in suricata.yaml, or Snort's inline -Q flag usage"),
        ("ID-7", "Management Access", "Management interface/console access is restricted and authenticated, not exposed on the production network", "General best practice", "netstat -tulpn  on the sensor's management interface, and review the firewall rules restricting access to it"),
        ("ID-8", "High Availability", "Fail-open vs fail-closed behavior for inline deployments is defined, documented and tested", "General best practice", "Review the inline bypass NIC configuration (bypass switch settings) and test behavior during a maintenance window"),
        ("ID-9", "Patch Management", "Sensor OS and engine (Snort/Suricata) kept on a supported, patched version", "General best practice", "suricata --build-info  or  snort -V  compared against the vendor's current release"),
        ("ID-10", "Tuning", "Baseline of normal traffic established; suppression/threshold rules used to reduce alert fatigue", "Vendor best practice", "Review suppress/threshold.conf entries and the alert-volume trend over time"),
        ("ID-11", "Correlation", "IDS/IPS alerts correlated with other telemetry (EDR, firewall logs) rather than reviewed in isolation", "General best practice", "Review SIEM correlation rules that reference the IDS/IPS as an alert source"),
    ],
    "PAM - CyberArk / BeyondTrust / Delinea (generic - no single CIS benchmark)": [
        ("PM-1", "Vaulting", "All privileged credentials (local admin, service, DB, cloud) are onboarded; no known unmanaged privileged accounts remain", "NIST SP 800-53 AC-2/AC-6", "Run the vendor's discovery/account-inventory report and compare it against the AD/local admin account list"),
        ("PM-2", "Password Rotation", "Automatic password rotation enabled per policy for vaulted accounts", "General PAM best practice", "Review the vault's rotation policy configuration and the last-rotated timestamp report"),
        ("PM-3", "Session Management", "Privileged sessions are brokered through the PAM solution with recording enabled for critical systems", "General PAM best practice", "Review the PSM/session-management policy configuration and sample a recorded session"),
        ("PM-4", "Access Control", "Just-in-time (JIT) or time-bound access used instead of standing privileged access where feasible", "Zero Trust / General PAM best practice", "Review the access-policy configuration for time-bound/ephemeral access grants"),
        ("PM-5", "MFA", "Multi-factor authentication required to check out credentials or start a privileged session", "General PAM best practice", "Attempt a credential checkout or session initiation and confirm an MFA prompt; review the auth policy config"),
        ("PM-6", "Least Privilege", "Access requests are approved via workflow (dual control) for high-risk systems", "General PAM best practice", "Review the approval-workflow configuration for high-risk safes/systems"),
        ("PM-7", "Auditing", "All vault access and admin actions on the PAM platform itself are logged and forwarded to SIEM", "General PAM best practice", "Review the vault's own audit log (e.g. CyberArk PVWA Reports, BeyondTrust Log/Audit) and confirm SIEM forwarding"),
        ("PM-8", "Break-glass", "Emergency/break-glass access procedure defined, tested, and monitored/alerted on use", "General PAM best practice", "Review the documented emergency-access SOP and check the alerting rule that triggers on its use"),
        ("PM-9", "Discovery", "Periodic automated discovery scans run to find newly created or orphaned privileged accounts outside the vault", "General PAM best practice", "Review the scheduled discovery job configuration and its most recent run report for orphaned accounts"),
        ("PM-10", "High Availability", "PAM vault/platform deployed in HA; backup of vault data encrypted and restore tested", "General PAM best practice", "Review the vault cluster/HA topology diagram and the most recent backup restore-test log"),
        ("PM-11", "Integration", "PAM integrated with SSO/IdP and ticketing so privileged access ties to a documented business justification", "General PAM best practice", "Review the IdP integration configuration and the ticket-linkage field in the access-request workflow"),
        ("PM-12", "Segregation", "PAM platform's own admin accounts are separate from the accounts it manages, with restricted vault console access", "General PAM best practice", "Review the PAM platform's own admin role assignments against the safes/permissions it manages"),
    ],
}

STATUS_OPTIONS = ["Not Reviewed", "Compliant", "Non-Compliant", "Compensating Control", "Not Applicable"]
STATUS_COLORS = {
    "Compliant": "#2ecc71",
    "Non-Compliant": "#e74c3c",
    "Compensating Control": "#f39c12",
    "Not Applicable": "#95a5a6",
    "Not Reviewed": "#bdc3c7",
}

# ---------------------------------------------------------------------------
# 2. SESSION STATE
# ---------------------------------------------------------------------------

if "responses" not in st.session_state:
    # key: item_id -> dict(status, notes, reviewer, date)
    st.session_state.responses = {}

if "client_name" not in st.session_state:
    st.session_state.client_name = ""

if "custom_checklists" not in st.session_state:
    # technologies drafted via the AI "add new technology" feature this session.
    # key: technology name -> list of (item_id, category, control, reference, audit_step)
    st.session_state.custom_checklists = {}

if "exec_summary" not in st.session_state:
    st.session_state.exec_summary = ""

# Built-in + any AI-drafted technologies, merged for use everywhere below
ALL_CHECKLISTS = {**CHECKLISTS, **st.session_state.custom_checklists}

# ---------------------------------------------------------------------------
# 3. SIDEBAR - client & technology setup, import/export
# ---------------------------------------------------------------------------

st.sidebar.markdown("### 🛡️ Sentinel GRC")
st.sidebar.markdown(f'<div class="library-badge">📚 {len(ALL_CHECKLISTS)} technologies in library</div>', unsafe_allow_html=True)
cloud_dot = "dot-on" if cloud_persistence_enabled() else "dot-off"
ai_reason = hf_diagnose()
ai_dot = "dot-off" if ai_reason else "dot-on"
st.sidebar.markdown(
    f'<div class="library-badge"><span class="status-dot {cloud_dot}"></span>Cloud sync</div>'
    f'<div class="library-badge"><span class="status-dot {ai_dot}"></span>AI agent</div>',
    unsafe_allow_html=True,
)
if ai_reason:
    st.sidebar.caption(f"⚠️ AI agent off: {ai_reason}")
st.sidebar.title("⚙️ Engagement Setup")
st.session_state.client_name = st.sidebar.text_input("Client / Engagement name", value=st.session_state.client_name)
reviewer = st.sidebar.text_input("Reviewer name", value="")
tech = st.sidebar.selectbox("Technology to assess", list(ALL_CHECKLISTS.keys()))

st.sidebar.markdown("---")
st.sidebar.subheader("🤖 Add a new technology with AI")
if not ai_enabled():
    st.sidebar.caption(f"⚠️ {hf_diagnose()}")
else:
    new_tech_name = st.sidebar.text_input("Technology name", placeholder="e.g. Kong API Gateway", key="new_tech_input")
    if st.sidebar.button("🤖 Draft checklist with AI", use_container_width=True):
        if not new_tech_name.strip():
            st.sidebar.warning("Enter a technology name first.")
        else:
            with st.sidebar.status("Agent is drafting and self-checking...", expanded=False):
                draft_prompt = f"""You are a senior GRC/OffSec consultant. Draft a hardening/configuration
checklist for: {new_tech_name}

Return ONLY valid JSON, no markdown fences, no commentary, in this exact shape:
{{
  "has_cis_benchmark": true/false,
  "items": [
    {{"category": "...", "control": "...", "reference": "CIS x.x or framework name", "audit_step": "specific command/GUI path to verify this", "severity": "Critical|High|Medium|Low"}}
  ]
}}
Produce 10-15 items across a sensible spread of categories (access control,
logging, network/encryption, patching, hardening). If no formal CIS Benchmark
exists for this technology, set has_cis_benchmark to false and base items on
the vendor's own security guide or general best practice (NIST/OWASP),
labeling the reference field accordingly. Be specific and technical."""
                parsed, attempts_used, raw = agentic_json_call(
                    draft_prompt,
                    system="You output only raw JSON, never prose or markdown fences.",
                )

            if parsed is None:
                if raw.startswith("⚠️"):
                    st.sidebar.error(raw)
                else:
                    st.sidebar.error(f"Agent gave up after {attempts_used} self-correction attempts — still not valid JSON.")
                    with st.sidebar.expander("Last raw output"):
                        st.code(raw)
            else:
                try:
                    prefix = "".join(w[0] for w in new_tech_name.upper().split())[:3] or "CU"
                    items = []
                    for idx, it in enumerate(parsed.get("items", []), start=1):
                        items.append((
                            f"{prefix}-{idx}",
                            it.get("category", "General"),
                            it.get("control", ""),
                            it.get("reference", "AI-drafted, unverified"),
                            it.get("audit_step", "Manually determine verification method"),
                        ))
                    label = f"{new_tech_name} (AI-drafted — review before client use)"
                    st.session_state.custom_checklists[label] = items
                    attempt_note = " on the first try" if attempts_used == 1 else f" after {attempts_used} self-correction attempts"
                    st.sidebar.success(f"Agent drafted {len(items)} controls for {new_tech_name}{attempt_note}. Select it above.")
                    st.rerun()
                except (KeyError, AttributeError) as e:
                    st.sidebar.error(f"Parsed JSON didn't match the expected checklist shape: {e}")
                    with st.sidebar.expander("Raw AI output"):
                        st.code(raw)

st.sidebar.markdown("---")
st.sidebar.subheader("💾 Save / Load Progress")

if cloud_persistence_enabled():
    st.sidebar.success("☁️ Cloud persistence connected")
    cc1, cc2 = st.sidebar.columns(2)
    if cc1.button("☁️ Save to cloud", use_container_width=True):
        save_to_cloud(st.session_state.client_name, tech, st.session_state.responses)
        st.sidebar.success("Saved.")
    if cc2.button("☁️ Load from cloud", use_container_width=True):
        loaded = load_from_cloud(st.session_state.client_name, tech)
        if loaded:
            st.session_state.responses.update(loaded)
            st.sidebar.success(f"Loaded {len(loaded)} items.")
        else:
            st.sidebar.warning("No saved data found for this client/technology.")
else:
    st.sidebar.caption(
        "☁️ Cloud persistence not configured — using manual CSV export/import "
        "below. See README.md to add Google Sheets persistence."
    )

# Export
if st.session_state.responses:
    export_rows = []
    for item_id, data in st.session_state.responses.items():
        export_rows.append({"item_id": item_id, **data})
    export_df = pd.DataFrame(export_rows)
    csv_buf = io.StringIO()
    export_df.to_csv(csv_buf, index=False)
    st.sidebar.download_button(
        "⬇️ Export progress (CSV)",
        data=csv_buf.getvalue(),
        file_name=f"{(st.session_state.client_name or 'client').replace(' ', '_')}_compliance_progress.csv",
        mime="text/csv",
    )

uploaded = st.sidebar.file_uploader("⬆️ Resume from CSV", type=["csv"])
if uploaded is not None:
    resume_df = pd.read_csv(uploaded)
    for _, row in resume_df.iterrows():
        st.session_state.responses[row["item_id"]] = {
            "status": row.get("status", "Not Reviewed"),
            "notes": row.get("notes", "") if pd.notna(row.get("notes", "")) else "",
            "reviewer": row.get("reviewer", "") if pd.notna(row.get("reviewer", "")) else "",
            "date": row.get("date", "") if pd.notna(row.get("date", "")) else "",
            "category": row.get("category", ""),
            "control": row.get("control", ""),
            "reference": row.get("reference", ""),
            "audit_step": row.get("audit_step", "") if pd.notna(row.get("audit_step", "")) else "",
        }
    st.sidebar.success(f"Loaded {len(resume_df)} saved responses.")

st.sidebar.markdown("---")
st.sidebar.caption(
    "Checklists are built from CIS Benchmark categories where a published "
    "benchmark exists for the technology. Where no formal CIS Benchmark "
    "applies (e.g. custom web apps), a generic best-practice checklist is "
    "used instead. Always validate against the current official benchmark "
    "PDF for the exact build/version in scope before sign-off."
)

# ---------------------------------------------------------------------------
# 4. MAIN - Tabs: Checklist / Dashboard
# ---------------------------------------------------------------------------

st.markdown(f"""
<div class="hero-banner">
    <p class="hero-title">🛡️ Sentinel GRC</p>
    <p class="hero-sub">Enterprise Security Configuration &amp; Compliance Platform</p>
    <p class="hero-meta">
        <b>Client:</b> {st.session_state.client_name or '—'} &nbsp;|&nbsp;
        <b>Technology:</b> {tech} &nbsp;|&nbsp;
        <b>Reviewer:</b> {reviewer or '—'} &nbsp;|&nbsp;
        <b>Library:</b> {len(ALL_CHECKLISTS)} technologies, CIS/vendor-benchmark aligned
    </p>
</div>
""", unsafe_allow_html=True)

tab_checklist, tab_dashboard, tab_engagements = st.tabs(
    ["📋 Checklist", "📊 Compliance Dashboard", "📁 Engagements"]
)

items = ALL_CHECKLISTS[tech]
categories = sorted(set(c for _, c, _, _, _ in items))

SEVERITY_PILL_CLASS = {"Critical": "sev-critical", "High": "sev-high", "Medium": "sev-medium", "Low": "sev-low"}

with tab_checklist:
    st.subheader("Configuration Checklist")
    filter_status = st.multiselect("Filter by status", STATUS_OPTIONS, default=[])
    filter_severity = st.multiselect("Filter by severity", list(SEVERITY_WEIGHTS.keys()), default=[])
    search = st.text_input("Search controls", "")

    for cat in categories:
        cat_items = [it for it in items if it[1] == cat]
        if search:
            cat_items = [it for it in cat_items if search.lower() in it[2].lower()]
        if not cat_items:
            continue

        with st.expander(f"**{cat}**  ({len(cat_items)} controls)", expanded=False):
            for item_id, category, control, ref, audit_step in cat_items:
                classification = classify_control(category, control)
                severity = classification["severity"]
                frameworks = classification["frameworks"]

                if filter_severity and severity not in filter_severity:
                    continue

                existing = st.session_state.responses.get(item_id, {})
                current_status = existing.get("status", "Not Reviewed")
                if filter_status and current_status not in filter_status:
                    continue

                col1, col2 = st.columns([3, 1])
                with col1:
                    pill_class = SEVERITY_PILL_CLASS.get(severity, "sev-low")
                    fw_line = " · ".join(f"<b>{k}:</b> {v}" for k, v in frameworks.items())
                    st.markdown(f"""
<div class="control-card">
    <span class="control-id">{item_id}</span> &nbsp;<span class="status-pill {pill_class}">{severity}</span>
    <div style="margin-top:6px; color:#e2e8f0; font-size:14.5px;">{control}</div>
    <div class="control-ref">Reference: {ref}</div>
    <div class="control-step">🔍 {audit_step}</div>
    <div class="fw-line">{fw_line}</div>
</div>
""", unsafe_allow_html=True)
                with col2:
                    status = st.selectbox(
                        "Status", STATUS_OPTIONS,
                        index=STATUS_OPTIONS.index(current_status),
                        key=f"status_{item_id}",
                        label_visibility="collapsed",
                    )
                notes = st.text_area(
                    "Evidence / notes", value=existing.get("notes", ""),
                    key=f"notes_{item_id}", height=60,
                    placeholder="Evidence collected, deviation justification, remediation owner/date...",
                )

                if ai_enabled():
                    if st.button("🤖 Suggest verdict from evidence above", key=f"ai_verdict_{item_id}"):
                        if not notes.strip():
                            st.warning("Paste command output or evidence in the notes field first.")
                        else:
                            with st.spinner("Analyzing evidence..."):
                                verdict_prompt = f"""Control: {control}
Audit step: {audit_step}
Evidence/output provided by the reviewer:
---
{notes}
---
Based only on this evidence, suggest one status: Compliant, Non-Compliant, or
Compensating Control. Give a one-sentence reason. Format your reply exactly as:
STATUS: <status>
REASON: <one sentence>"""
                            suggestion = ask_ai(verdict_prompt, system="You are a precise security auditor. Never invent evidence not given to you.")
                            if suggestion.startswith("⚠️"):
                                st.error(suggestion)
                            else:
                                st.info(f"🤖 AI suggestion (confirm before applying): {suggestion}")

                st.session_state.responses[item_id] = {
                    "status": status,
                    "notes": notes,
                    "category": category,
                    "control": control,
                    "reference": ref,
                    "audit_step": audit_step,
                    "severity": severity,
                    "frameworks": frameworks,
                    "reviewer": reviewer,
                    "date": datetime.now().strftime("%Y-%m-%d"),
                }
                st.markdown("---")

with tab_dashboard:
    st.subheader("Compliance Dashboard")

    all_ids = [it[0] for it in items]
    item_lookup = {it[0]: it for it in items}
    recorded = {}
    for i in all_ids:
        r = st.session_state.responses.get(i)
        if r:
            recorded[i] = r
        else:
            _, cat, ctrl, _, _ = item_lookup[i]
            recorded[i] = {"status": "Not Reviewed", "category": cat, "severity": classify_control(cat, ctrl)["severity"]}

    df = pd.DataFrame([
        {
            "item_id": i,
            "category": recorded[i].get("category", ""),
            "status": recorded[i].get("status", "Not Reviewed"),
            "severity": recorded[i].get("severity") or classify_control(recorded[i].get("category", ""), "")["severity"],
        }
        for i in all_ids
    ])
    df["weight"] = df["severity"].map(SEVERITY_WEIGHTS).fillna(1)

    total = len(df)
    compliant = (df["status"] == "Compliant").sum()
    noncompliant = (df["status"] == "Non-Compliant").sum()
    compensating = (df["status"] == "Compensating Control").sum()
    na = (df["status"] == "Not Applicable").sum()
    not_reviewed = (df["status"] == "Not Reviewed").sum()
    applicable = total - na
    compliance_pct = round(((compliant + compensating) / applicable) * 100, 1) if applicable else 0.0

    applicable_df = df[df["status"] != "Not Applicable"]
    good_weight = applicable_df[applicable_df["status"].isin(["Compliant", "Compensating Control"])]["weight"].sum()
    total_weight = applicable_df["weight"].sum()
    weighted_compliance_pct = round((good_weight / total_weight) * 100, 1) if total_weight else 0.0

    m0, m1, m2, m3, m4, m5 = st.columns(6)
    m0.metric("Risk-Weighted Compliance", f"{weighted_compliance_pct}%",
              help="Weighted by severity: Critical×4, High×3, Medium×2, Low×1. A Critical finding moves this more than a Low one.")
    m1.metric("Raw Compliance", f"{compliance_pct}%", help="Every applicable control counted equally.")
    m2.metric("Compliant", int(compliant))
    m3.metric("Non-Compliant", int(noncompliant))
    m4.metric("Compensating Control", int(compensating))
    m5.metric("Not Reviewed", int(not_reviewed))

    crit_open = df[(df["status"] == "Non-Compliant") & (df["severity"] == "Critical")]
    if len(crit_open):
        st.error(f"🔴 {len(crit_open)} Critical-severity control(s) are currently Non-Compliant — prioritize these first.")

    st.markdown("---")

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**Status breakdown**")
        status_counts = df["status"].value_counts().reindex(STATUS_OPTIONS, fill_value=0).reset_index()
        status_counts.columns = ["status", "count"]
        fig_pie = px.pie(
            status_counts, names="status", values="count",
            color="status", color_discrete_map=STATUS_COLORS, hole=0.45,
        )
        fig_pie.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            font_color="#cbd5e1", legend_font_color="#cbd5e1",
            margin=dict(t=10, b=10, l=10, r=10),
        )
        st.plotly_chart(fig_pie, use_container_width=True)

    with c2:
        st.markdown("**Compliance % by category**")
        cat_stats = []
        for cat in categories:
            cat_df = df[df["category"] == cat]
            cat_applicable = cat_df[cat_df["status"] != "Not Applicable"]
            cat_good = cat_applicable[cat_applicable["status"].isin(["Compliant", "Compensating Control"])]
            pct = round((len(cat_good) / len(cat_applicable)) * 100, 1) if len(cat_applicable) else 0.0
            cat_stats.append({"category": cat, "compliance_pct": pct})
        cat_df_stats = pd.DataFrame(cat_stats).sort_values("compliance_pct")
        fig_bar = px.bar(
            cat_df_stats, x="compliance_pct", y="category", orientation="h",
            range_x=[0, 100], text="compliance_pct",
            color="compliance_pct", color_continuous_scale=["#e74c3c", "#f39c12", "#2ecc71"],
        )
        fig_bar.update_traces(texttemplate="%{text}%", textposition="outside")
        fig_bar.update_layout(
            coloraxis_showscale=False,
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            font_color="#cbd5e1",
            xaxis=dict(gridcolor="#1e293b"), yaxis=dict(gridcolor="#1e293b"),
            margin=dict(t=10, b=10, l=10, r=10),
        )
        st.plotly_chart(fig_bar, use_container_width=True)

    st.markdown("---")
    st.markdown("**⚠️ Open Non-Compliant Items (remediation tracker)**")
    open_items = []
    for i in all_ids:
        r = st.session_state.responses.get(i)
        if r and r.get("status") == "Non-Compliant":
            open_items.append({
                "Item": i, "Category": r.get("category"), "Severity": r.get("severity"),
                "Control": r.get("control"), "Notes": r.get("notes"),
                "Reference": r.get("reference"), "Audit Step": r.get("audit_step"),
            })
    if open_items:
        open_df = pd.DataFrame(open_items)
        sev_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
        open_df["_sort"] = open_df["Severity"].map(sev_order).fillna(4)
        open_df = open_df.sort_values("_sort").drop(columns="_sort")
        st.dataframe(open_df, use_container_width=True, hide_index=True)
    else:
        st.success("No open non-compliant items recorded yet for this technology.")

    st.markdown("---")
    st.markdown("**🧭 Framework cross-mapping**")
    st.caption("Rule-based mapping from control category to framework clause — a starting point for cross-framework reporting, confirm against the actual framework text before client delivery.")
    fw_rows = []
    for i in all_ids:
        r = st.session_state.responses.get(i) or {}
        fw = r.get("frameworks") or classify_control(r.get("category", ""), r.get("control", ""))["frameworks"]
        fw_rows.append({"Item": i, "Category": r.get("category", ""), **fw})
    st.dataframe(pd.DataFrame(fw_rows).drop_duplicates(subset="Category"), use_container_width=True, hide_index=True)

    st.markdown("---")
    st.markdown("**🤖 AI-assisted reporting**")
    if not ai_enabled():
        st.caption(f"⚠️ {hf_diagnose()}")
    else:
        ai_col1, ai_col2 = st.columns(2)
        with ai_col1:
            if st.button("🤖 Generate Executive Summary", use_container_width=True):
                with st.spinner("Drafting summary..."):
                    top_findings = open_df.head(5).to_dict("records") if open_items else []
                    summary_prompt = f"""Write a concise executive summary (150-250 words) for a security
configuration assessment.
Client: {st.session_state.client_name or 'the client'}
Technology assessed: {tech}
Risk-weighted compliance score: {weighted_compliance_pct}%
Total controls: {total}, Non-Compliant: {int(noncompliant)}, Compensating Control: {int(compensating)}, Not Reviewed: {int(not_reviewed)}
Top open findings (highest severity first): {top_findings}

Write in a professional GRC consulting tone, third person, suitable to paste
directly into a client report. Cover: overall posture, the most urgent risks,
and a one-line recommended next step. Do not invent findings not listed above."""
                    summary = ask_ai(summary_prompt, system="You are a senior GRC consultant writing a client-facing report section.")
                    if summary.startswith("⚠️"):
                        st.error(summary)
                    else:
                        st.session_state.exec_summary = summary
        with ai_col2:
            st.caption("Generated summary is included in the DOCX report export below, and can be edited before export.")

        if st.session_state.exec_summary:
            st.session_state.exec_summary = st.text_area(
                "Executive summary (editable — this is what gets exported)",
                value=st.session_state.exec_summary, height=180,
            )

        st.markdown("&nbsp;")
        nl_question = st.text_input("Ask a question about this assessment", placeholder="e.g. Which encryption-related controls are still open?")
        if st.button("Ask") and nl_question.strip():
            with st.spinner("Thinking..."):
                context_rows = [
                    {"item_id": i, "category": r.get("category"), "status": r.get("status"),
                     "severity": r.get("severity"), "control": r.get("control"), "notes": r.get("notes")}
                    for i, r in st.session_state.responses.items()
                ]
                qa_prompt = f"""Assessment data for {tech} ({st.session_state.client_name or 'client'}):
{json.dumps(context_rows)[:6000]}

Question: {nl_question}

Answer using only the data above. If the data doesn't cover the question, say so."""
                answer = ask_ai(qa_prompt, system="You answer strictly from the provided assessment data, citing item IDs where relevant.")
                if answer.startswith("⚠️"):
                    st.error(answer)
                else:
                    st.info(answer)

    st.markdown("---")
    st.markdown("**📄 Export**")
    exp_col1, exp_col2 = st.columns(2)
    with exp_col1:
        if st.session_state.responses:
            full_export = []
            for item_id, data in st.session_state.responses.items():
                full_export.append({"item_id": item_id, **data})
            full_df = pd.DataFrame(full_export)
            csv_buf2 = io.StringIO()
            full_df.to_csv(csv_buf2, index=False)
            st.download_button(
                "⬇️ Download full compliance report (CSV)",
                data=csv_buf2.getvalue(),
                file_name=f"{(st.session_state.client_name or 'client').replace(' ', '_')}_{tech.split(' ')[0]}_compliance_report.csv",
                mime="text/csv",
                use_container_width=True,
            )

    with exp_col2:
        if st.session_state.responses:
            docx_bytes = build_docx_report(
                client_name=st.session_state.client_name or "Client",
                technology=tech,
                reviewer_name=reviewer,
                weighted_pct=weighted_compliance_pct,
                raw_pct=compliance_pct,
                totals={"total": total, "compliant": int(compliant), "noncompliant": int(noncompliant),
                        "compensating": int(compensating), "not_reviewed": int(not_reviewed)},
                open_items=open_items,
                exec_summary=st.session_state.exec_summary,
                all_responses=st.session_state.responses,
            )
            st.download_button(
                "⬇️ Download client report (DOCX)",
                data=docx_bytes,
                file_name=f"{(st.session_state.client_name or 'client').replace(' ', '_')}_{tech.split(' ')[0]}_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

with tab_engagements:
    st.subheader("Engagements")
    if not cloud_persistence_enabled():
        st.info(
            "Engagement listing requires cloud persistence (Google Sheets) to be configured, "
            "since each engagement is stored as a tab in the shared sheet. See README.md."
        )
    else:
        try:
            gc = get_gsheet_client()
            sh = gc.open_by_key(st.secrets["gsheets"]["sheet_id"])
            worksheets = sh.worksheets()
            rows = []
            for ws in worksheets:
                name = ws.title
                if "__" in name:
                    client_part, tech_part = name.split("__", 1)
                else:
                    client_part, tech_part = name, ""
                try:
                    row_count = max(ws.row_count - 1, 0)
                    records = ws.get_all_records()
                    nc = sum(1 for r in records if r.get("status") == "Non-Compliant")
                    good = sum(1 for r in records if r.get("status") in ("Compliant", "Compensating Control"))
                    applicable = sum(1 for r in records if r.get("status") != "Not Applicable")
                    pct = round((good / applicable) * 100, 1) if applicable else None
                except Exception:
                    records, nc, pct = [], 0, None
                rows.append({
                    "Client": client_part, "Technology": tech_part,
                    "Controls Recorded": len(records), "Open Non-Compliant": nc,
                    "Compliance %": pct,
                })

            if rows:
                portfolio_df = pd.DataFrame(rows)
                scored = portfolio_df[portfolio_df["Compliance %"].notna()]

                p1, p2, p3 = st.columns(3)
                p1.metric("Active Engagements", len(portfolio_df))
                p2.metric("Total Open Non-Compliant", int(portfolio_df["Open Non-Compliant"].sum()))
                p3.metric("Portfolio Avg. Compliance", f"{round(scored['Compliance %'].mean(), 1)}%" if len(scored) else "—")

                if len(scored):
                    st.markdown("**Compliance by engagement**")
                    chart_df = scored.copy()
                    chart_df["Engagement"] = chart_df["Client"] + " — " + chart_df["Technology"]
                    chart_df = chart_df.sort_values("Compliance %")
                    fig_portfolio = px.bar(
                        chart_df, x="Compliance %", y="Engagement", orientation="h",
                        range_x=[0, 100], text="Compliance %",
                        color="Compliance %", color_continuous_scale=["#e74c3c", "#f39c12", "#2ecc71"],
                    )
                    fig_portfolio.update_traces(texttemplate="%{text}%", textposition="outside")
                    fig_portfolio.update_layout(
                        coloraxis_showscale=False,
                        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                        font_color="#cbd5e1",
                        xaxis=dict(gridcolor="#1e293b"), yaxis=dict(gridcolor="#1e293b"),
                        margin=dict(t=10, b=10, l=10, r=10),
                    )
                    st.plotly_chart(fig_portfolio, use_container_width=True)

                st.markdown("**All engagements**")
                st.dataframe(portfolio_df, use_container_width=True, hide_index=True)
                st.caption("Select a client + technology in the sidebar and click '☁️ Load from cloud' to resume any of these engagements.")
            else:
                st.info("No engagements saved yet. Save progress from the sidebar to see it listed here.")
        except Exception as e:
            st.error(f"Could not list engagements: {e}")
